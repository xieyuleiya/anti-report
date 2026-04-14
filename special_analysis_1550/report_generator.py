# 这里可根据需要实现Word报告生成等功能

import os
import shutil
from datetime import datetime
try:
    from .utils import log
    from .charts import create_critical_sensors_chart, create_trend_chart
except ImportError:
    from utils import log
    from charts import create_critical_sensors_chart, create_trend_chart

# ==================== 配置常量 ====================
try:
    from docx import Document
    from docx.shared import Inches, Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.shared import OxmlElement, qn
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    log("缺少 python-docx 库，请安装: pip install python-docx", 'warn')

# 文件路径配置（由主程序注入；可从环境变量 REPORT_TEMPLATE_PATH 读取；为空时使用默认空文档）
TEMPLATE_PATH = None
# 当前报告输出目录（用于存放临时图片等过程文件）
REPORT_OUTPUT_DIR = None


# from typing import Union

# def set_template_path(path: Union[str, None]):
def set_template_path(path=None):
    """供主程序设置报告模板路径。"""
    global TEMPLATE_PATH
    TEMPLATE_PATH = path

# 样式配置
STYLES = {
    'TABLE_INNER_TEXT': '表内文字',
    'PICTURE': '图格式',
    'CAPTION': '题注',
    'CAPTION_FALLBACK': 'Caption',
    'TABLE': 'Light Grid Accent 1'
}

# 报告类型配置
REPORT_TYPES = {
    'MONTH': {
        'filename': '月报.docx',
        'monitoring_title': '监测结果',
        'table_headers': ['测点编号', '横向倾角', '纵向倾角', '测点编号', '横向倾角', '纵向倾角'],
        'table_cols': 6
    },
    'QUARTER': {
        'filename': '季报.docx',
        'monitoring_title': '监测结果',
        'table_headers': ['测点', '横向倾角\n(°)', '横向趋势', '横向最大\n(°)', '横向最小\n(°)', '横向变化\n(°)', 
                         '纵向倾角\n(°)', '纵向趋势', '纵向最大\n(°)', '纵向最小\n(°)', '纵向变化\n(°)'],
        'table_cols': 11
    }
}

# 错误消息
ERROR_MESSAGES = {
    'DOCX_NOT_AVAILABLE': "❌ 无法生成Word报告，请安装 python-docx 库",
    'TEMPLATE_NOT_FOUND': "❌ 模板文件不存在: {}",
    'TEMPLATE_LOAD_FAILED': "❌ 加载模板文件失败: {}",
    'TEMPLATE_FALLBACK': "❌ 无法加载模板文档，将使用默认文档。",
    'STYLE_NOT_FOUND': "⚠️ 样式 '{}' 不存在，图片将使用默认段落样式。",
    'MONITORING_TITLE_NOT_FOUND': "⚠️ 未在模板中找到'{}'标题。",
    'CONCLUSION_TITLE_NOT_FOUND': "⚠️ 未在模板中找到'结论'标题。",
    'CHART_GENERATION_FAILED': "❌ 图表生成失败 {}",
    'CHART_INSERT_FAILED': "❌ 插入图表失败 {}: {}",
    'SAVE_FAILED': "❌ 保存{}失败: {}",
    'BORDER_ADD_FAILED': "⚠️ 添加表格边框失败: {}"
}

# 成功消息
SUCCESS_MESSAGES = {
    'REPORT_SAVED': "✅ {}已保存至: {}",
    'CHART_GENERATED': "✅ 图表生成成功，正在插入文档...",
    'CHART_INSERTED': "✅ 图表已成功插入文档",
    'TEMPLATE_LOADED': "✅ 成功加载模板文件: {}"
}

# 信息消息
INFO_MESSAGES = {
    'LOCATING_TITLE': "✅ 定位到'{}'标题，正在插入内容...",
    'GENERATING_CHART': "🎨 开始为 {} 生成持续关注测点图表...",
    'ADDING_SENSOR_DATA': "  ✅ 添加测点数据: {}",
    'GENERATING_CHART_FILE': "  📊 正在生成图表: {}"
}

# 警告消息
WARNING_MESSAGES = {
    'NO_CRITICAL_SENSORS': "⚠️ {} 没有持续关注测点，跳过图表生成",
    'NO_BRIDGE_DATA': "⚠️ {} 缺少测点数据，跳过图表生成"
}

# ==================== 工具函数 ====================

def load_template_document():
    """加载模板文档"""
    if not DOCX_AVAILABLE:
        return None
    
    if not TEMPLATE_PATH:
        log(ERROR_MESSAGES['TEMPLATE_NOT_FOUND'].format("未指定模板路径"), 'error')
        return None
    
    if not os.path.exists(TEMPLATE_PATH):
        log(ERROR_MESSAGES['TEMPLATE_NOT_FOUND'].format(TEMPLATE_PATH), 'error')
        return None
    
    try:
        doc = Document(TEMPLATE_PATH)
        return doc
    except Exception as e:
        log(ERROR_MESSAGES['TEMPLATE_LOAD_FAILED'].format(e), 'error')
        return None

def _set_cell_format(cell, text, is_critical=False):
    """设置单元格格式"""
    if not DOCX_AVAILABLE:
        return
    cell.text = ''
    paragraph = cell.paragraphs[0]
    paragraph.style = STYLES['TABLE_INNER_TEXT']
    run = paragraph.add_run(text)
    if is_critical or text == "持续关注":
        run.font.color.rgb = RGBColor(255, 0, 0)  # 红色
        run.bold = True
    elif text == "离线":
        run.font.color.rgb = RGBColor(255, 0, 0)  # 红色
        run.bold = True
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

def set_table_inner_text_style(table):
    """将表格内所有单元格的段落样式设置为'表内文字'"""
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                para.style = STYLES['TABLE_INNER_TEXT']

def add_table_border(table):
    """添加表格边框"""
    if not DOCX_AVAILABLE:
        return
    try:
        from docx.oxml import parse_xml
        from docx.oxml.ns import nsdecls
        tbl = table._tbl
        tblPr = tbl.tblPr
        borders_xml = (
            '<w:tblBorders %s>'
            '  <w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
            '  <w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
            '  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
            '  <w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
            '  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
            '  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
            '</w:tblBorders>' % nsdecls('w')
        )
        tblBorders = parse_xml(borders_xml)
        tblPr.append(tblBorders)
    except Exception as e:
        print(ERROR_MESSAGES['BORDER_ADD_FAILED'].format(e))

def add_field(paragraph, field_code):
    """在段落中插入一个域代码"""
    run = paragraph.add_run()
    r = run._r
    
    # 开始符
    fldChar_begin = OxmlElement('w:fldChar')
    fldChar_begin.set(qn('w:fldCharType'), 'begin')
    r.append(fldChar_begin)
    
    # 域代码指令
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = field_code
    r.append(instrText)
    
    # 分隔符
    fldChar_separate = OxmlElement('w:fldChar')
    fldChar_separate.set(qn('w:fldCharType'), 'separate')
    r.append(fldChar_separate)
    
    # 结束符
    fldChar_end = OxmlElement('w:fldChar')
    fldChar_end.set(qn('w:fldCharType'), 'end')
    r.append(fldChar_end)

def add_caption(doc, type_name, text):
    """为图或表添加带章节号的自动编号标题"""
    style_to_use = None
    if STYLES['CAPTION'] in doc.styles:
        style_to_use = STYLES['CAPTION']
    elif STYLES['CAPTION_FALLBACK'] in doc.styles:
        style_to_use = STYLES['CAPTION_FALLBACK']
    
    p = doc.add_paragraph(style=style_to_use)
    p.add_run(f'{type_name} ')
    add_field(p, r'STYLEREF 1 \s')
    p.add_run('-')
    add_field(p, f'SEQ {type_name} \\* ARABIC \\s 1')
    p.add_run(f' {text}')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return p

# ==================== 数据处理函数 ====================

def get_critical_sensors(sensors_data):
    """获取持续关注测点列表（包含原“重点关注”与“持续关注”两档）"""
    return [
        sid for sid, res in sensors_data.items()
        if (
            res != "OFFLINE" and (
                res['horizontal_angle_trend']['trend_strength'] in ('持续关注', '重点关注') or
                res['vertical_angle_trend']['trend_strength'] in ('持续关注', '重点关注')
            )
        )
    ]

def get_bridge_critical_summary(all_trend_results):
    """获取所有桥梁的重点关注测点汇总"""
    bridge_critical = {}
    for bridge_name, sensors in all_trend_results.items():
        critical_points = get_critical_sensors(sensors)
        if critical_points:
            bridge_critical[bridge_name] = critical_points
    return bridge_critical

def prepare_critical_sensors_data(bridge_name, critical_sensors, all_bridges_data):
    """准备重点关注测点数据"""
    if bridge_name not in all_bridges_data:
        return {}
    
    critical_sensors_data = {}
    for sensor_id in critical_sensors:
        sensor_data = all_bridges_data[bridge_name].get(sensor_id)
        if sensor_data is not None and not (isinstance(sensor_data, str) and sensor_data == "OFFLINE"):
            critical_sensors_data[sensor_id] = sensor_data
            log(INFO_MESSAGES['ADDING_SENSOR_DATA'].format(sensor_id), 'info')
    
    return critical_sensors_data

# ==================== 文档操作函数 ====================

def find_paragraph_by_text(doc, text, occurrence=1):
    """查找包含指定文本的段落"""
    found_count = 0
    for p in doc.paragraphs:
        if text in p.text and p.style.name.startswith('Heading'):
            found_count += 1
            if found_count == occurrence:
                return p
    return None

def insert_element_after(element, new_element):
    """在指定元素后插入新元素"""
    # 处理不同类型的元素
    if hasattr(new_element, '_p'):
        # 段落、标题等元素
        element_to_insert = new_element._p
    elif hasattr(new_element, '_tbl'):
        # 表格元素
        element_to_insert = new_element._tbl
    else:
        # 其他元素
        element_to_insert = new_element
    
    element.addnext(element_to_insert)
    return element_to_insert

def create_bridge_heading(doc, bridge_name, cursor_element):
    """创建桥梁标题"""
    h3 = doc.add_heading(bridge_name, level=3)
    return insert_element_after(cursor_element, h3)

def create_intro_paragraph(doc, bridge_name, critical_sensors, cursor_element):
    """创建介绍段落"""
    base_text = f"{bridge_name}测点趋势强度明细如下表所示。"
    if critical_sensors:
        critical_text = '、'.join([f'{sid}测点' for sid in critical_sensors])
        text = f"{base_text}其中{critical_text}变化较明显，时序变化规律如下图所示。"
    else:
        text = base_text
    
    intro_p = doc.add_paragraph(text)
    return insert_element_after(cursor_element, intro_p)

def create_table_caption(doc, bridge_name, cursor_element):
    """创建表格标题"""
    caption_p = add_caption(doc, '表', f'{bridge_name}测点趋势强度明细')
    return insert_element_after(cursor_element, caption_p)

def create_picture_caption(doc, bridge_name, cursor_element):
    """创建图片标题"""
    caption_pic_p = add_caption(doc, '图', f'{bridge_name}持续关注测点趋势图')
    return insert_element_after(cursor_element, caption_pic_p)

# ==================== 表格创建函数 ====================

def create_month_table(doc, sensors_data):
    """创建月报表格"""
    sensor_ids = sorted(sensors_data.keys())
    n = len(sensor_ids)
    rows = (n + 1) // 2
    
    table = doc.add_table(rows=rows + 1, cols=REPORT_TYPES['MONTH']['table_cols'])
    table.style = STYLES['TABLE']
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # 设置表头
    hdr_cells = table.rows[0].cells
    headers = REPORT_TYPES['MONTH']['table_headers']
    for i, header in enumerate(headers):
        run = hdr_cells[i].paragraphs[0].add_run(header)
        run.bold = True
    
    # 填充数据
    for i in range(rows):
        for col_group in range(2):
            idx = i * 2 + col_group
            if idx < n:
                sensor_id = sensor_ids[idx]
                result = sensors_data[sensor_id]
                row_cells = table.rows[i + 1].cells
                row_cells[col_group * 3].text = sensor_id
                
                # 处理离线状态
                if result == "OFFLINE":
                    _set_cell_format(row_cells[col_group * 3 + 1], "离线", False)
                    _set_cell_format(row_cells[col_group * 3 + 2], "离线", False)
                else:
                    # 处理正常状态
                    h_strength = result['horizontal_angle_trend']['trend_strength']
                    v_strength = result['vertical_angle_trend']['trend_strength']
                    _set_cell_format(row_cells[col_group * 3 + 1], h_strength, h_strength == '持续关注')
                    _set_cell_format(row_cells[col_group * 3 + 2], v_strength, v_strength == '持续关注')
    
    add_table_border(table)
    set_table_inner_text_style(table)
    return table

def create_quarter_table(doc, sensors_data):
    """创建季报表格"""
    sensor_ids = sorted(sensors_data.keys())
    
    table = doc.add_table(rows=1, cols=REPORT_TYPES['QUARTER']['table_cols'])
    table.style = STYLES['TABLE']
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # 设置表头
    hdr_cells = table.rows[0].cells
    headers = REPORT_TYPES['QUARTER']['table_headers']
    for i, header in enumerate(headers):
        run = hdr_cells[i].paragraphs[0].add_run(header)
        run.bold = True
    
    # 填充数据
    for sensor_id in sensor_ids:
        result = sensors_data[sensor_id]
        row_cells = table.add_row().cells
        
        row_cells[0].text = sensor_id
        
        # 处理离线状态
        if result == "OFFLINE":
            row_cells[1].text = "离线"
            _set_cell_format(row_cells[2], "离线", False)
            row_cells[3].text = "离线"
            row_cells[4].text = "离线"
            row_cells[5].text = "离线"
            row_cells[6].text = "离线"
            _set_cell_format(row_cells[7], "离线", False)
            row_cells[8].text = "离线"
            row_cells[9].text = "离线"
            row_cells[10].text = "离线"
        else:
            # 处理正常状态
            h_trend = result['horizontal_angle_trend']
            v_trend = result['vertical_angle_trend']
            row_cells[1].text = f'{h_trend["slope_per_month"]:.4f}'
            _set_cell_format(row_cells[2], h_trend.get('trend_strength', '正常'), h_trend.get('trend_strength') == '持续关注')
            row_cells[3].text = f'{h_trend["max_value"]:.4f}'
            row_cells[4].text = f'{h_trend["min_value"]:.4f}'
            row_cells[5].text = f'{h_trend["value_range"]:.4f}'
            row_cells[6].text = f'{v_trend["slope_per_month"]:.4f}'
            _set_cell_format(row_cells[7], v_trend.get('trend_strength', '正常'), v_trend.get('trend_strength') == '持续关注')
            row_cells[8].text = f'{v_trend["max_value"]:.4f}'
            row_cells[9].text = f'{v_trend["min_value"]:.4f}'
            row_cells[10].text = f'{v_trend["value_range"]:.4f}'
    
    add_table_border(table)
    set_table_inner_text_style(table)
    return table

# ==================== 图表处理函数 ====================

def insert_chart(doc, bridge_name, critical_sensors_data, cursor_element, pic_style):
    """插入图表"""
    if not critical_sensors_data:
        log(WARNING_MESSAGES['NO_CRITICAL_SENSORS'].format(bridge_name), 'warn')
        return cursor_element
    
    # 将临时图片保存到输出目录下的临时子目录，由调用方生成并传入占位符
    # 由于本函数无法直接获知输出目录，这里退化为当前目录下生成，
    # 具体由调用方的包装函数负责传入正确的路径（在 generate_month_report / generate_quarter_report 中修正）。
    temp_img_path = f"temp_{bridge_name}_critical.png"
    log(INFO_MESSAGES['GENERATING_CHART_FILE'].format(temp_img_path), 'info')
    
    chart_result = create_critical_sensors_chart(bridge_name, critical_sensors_data, temp_img_path)
    
    if not chart_result:
        log(ERROR_MESSAGES['CHART_GENERATION_FAILED'].format(bridge_name), 'error')
        return cursor_element
    
    log(SUCCESS_MESSAGES['CHART_GENERATED'], 'info')
    
    try:
        # 插入图片
        p_for_pic = doc.add_paragraph(style=pic_style)
        p_for_pic.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_for_pic.add_run().add_picture(temp_img_path, width=Cm(16))
        os.remove(temp_img_path)
        cursor_element = insert_element_after(cursor_element, p_for_pic)
        
        # 插入图片标题
        caption_p = doc.add_paragraph()
        caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cursor_element = create_picture_caption(doc, bridge_name, cursor_element)
    except Exception as e:
        log(ERROR_MESSAGES['CHART_INSERT_FAILED'].format(bridge_name, str(e)), 'error')
    
    return cursor_element

# ==================== 结论生成函数 ====================

def generate_conclusion(doc, all_trend_results, all_bridges_data, cursor_element):
    """生成结论段落"""
    bridge_critical = get_bridge_critical_summary(all_trend_results)
    
    # 收集离线测点信息（按桥聚合测点）
    offline_by_bridge = {}
    for bridge_name, sensors_data in all_bridges_data.items():
        for sensor_id, sensor_data in sensors_data.items():
            if isinstance(sensor_data, str) and sensor_data == "OFFLINE":
                offline_by_bridge.setdefault(bridge_name, []).append(sensor_id)
    
    conclusion_parts = []
    
    # 重点关注测点部分
    if bridge_critical:
        parts = []
        for bridge, points in bridge_critical.items():
            parts.append(f"{bridge}{'、'.join(points)}")
        summary = '，'.join(parts)
        conclusion_parts.append(f"干线公路结构监测系统建设攻坚方案中{summary}为本监测期重点关注测点，从上述数据分析结果中可以看出，其横向和纵向倾角随时间变化规律较明显，在监测期间呈现出明显的趋势性变化，需要持续跟踪监测其发展态势。")
    else:
        conclusion_parts.append("本期未发现需重点关注的桥梁测点。")
    
    # 离线测点部分（按桥聚合）
    if offline_by_bridge:
        offline_parts = []
        for bridge, points in offline_by_bridge.items():
            offline_parts.append(f"{bridge}{'、'.join(points)}")
        offline_summary = '，'.join(offline_parts)
        offline_text = f"{offline_summary}测点倾角仪出现故障，运维团队将尽快安排人员前往现场进行检修与恢复工作。"
        conclusion_parts.append(offline_text)
    
    # 合并所有结论内容
    conclusion_text = "".join(conclusion_parts)
    para = doc.add_paragraph(conclusion_text)
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return insert_element_after(cursor_element, para)

# ==================== 桥梁内容生成函数 ====================

def generate_bridge_content(doc, bridge_name, sensors_data, all_bridges_data, cursor_element, pic_style, report_type):
    """生成单个桥梁的内容"""
    # 创建桥梁标题
    cursor_element = create_bridge_heading(doc, bridge_name, cursor_element)
    
    # 获取重点关注测点
    critical_sensors = get_critical_sensors(sensors_data)
    
    # 创建介绍段落
    cursor_element = create_intro_paragraph(doc, bridge_name, critical_sensors, cursor_element)
    
    # 创建表格标题
    cursor_element = create_table_caption(doc, bridge_name, cursor_element)
    
    # 创建表格
    if report_type == 'MONTH':
        table = create_month_table(doc, sensors_data)
    else:
        table = create_quarter_table(doc, sensors_data)
    
    cursor_element = insert_element_after(cursor_element, table)
    
    # 插入图表（如果有重点关注测点）
    if critical_sensors:
        critical_sensors_data = prepare_critical_sensors_data(bridge_name, critical_sensors, all_bridges_data)
        # 使用输出目录下的临时文件夹保存临时图片
        temp_dir = os.path.join(REPORT_OUTPUT_DIR or os.getcwd(), "临时图片")
        os.makedirs(temp_dir, exist_ok=True)
        # 临时替换 insert_chart 内部默认路径：在调用前先创建图片，再插入并删除
        temp_img_path = os.path.join(temp_dir, f"temp_{bridge_name}_critical.png")
        log(INFO_MESSAGES['GENERATING_CHART_FILE'].format(temp_img_path), 'info')
        chart_result = create_critical_sensors_chart(bridge_name, critical_sensors_data, temp_img_path)
        if chart_result:
            log(SUCCESS_MESSAGES['CHART_GENERATED'], 'info')
            try:
                p_for_pic = doc.add_paragraph(style=pic_style)
                p_for_pic.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p_for_pic.add_run().add_picture(temp_img_path, width=Cm(16))
                os.remove(temp_img_path)
                cursor_element = insert_element_after(cursor_element, p_for_pic)
                cursor_element = create_picture_caption(doc, bridge_name, cursor_element)
            except Exception as e:
                log(ERROR_MESSAGES['CHART_INSERT_FAILED'].format(bridge_name, str(e)), 'error')
        else:
            log(ERROR_MESSAGES['CHART_GENERATION_FAILED'].format(bridge_name), 'error')
    else:
        log(WARNING_MESSAGES['NO_CRITICAL_SENSORS'].format(bridge_name), 'warn')
    
    return cursor_element

# ==================== 主报告生成函数 ====================

def generate_report_base(all_trend_results, all_bridges_data, output_dir, report_type='MONTH'):
    # 记录全局输出目录供内部函数使用
    global REPORT_OUTPUT_DIR
    REPORT_OUTPUT_DIR = output_dir
    if not DOCX_AVAILABLE:
        log(ERROR_MESSAGES['DOCX_NOT_AVAILABLE'])
        return None
    
    # 加载模板文档
    doc = load_template_document()
    if doc is None:
        log(ERROR_MESSAGES['TEMPLATE_FALLBACK'])
        doc = Document()
    
    # 检查图片样式
    pic_style = STYLES['PICTURE'] if STYLES['PICTURE'] in doc.styles else 'Normal'
    if pic_style == 'Normal':
        log(ERROR_MESSAGES['STYLE_NOT_FOUND'].format(STYLES['PICTURE']), 'warn')
    
    # 设置报告路径
    report_path = os.path.join(output_dir, REPORT_TYPES[report_type]['filename'])
    
    # 查找插入点
    monitoring_title = REPORT_TYPES[report_type]['monitoring_title']
    monitoring_results_p = find_paragraph_by_text(doc, monitoring_title, 2)
    conclusion_p = find_paragraph_by_text(doc, '结论', 1)
    
    # 生成监测结果内容
    if monitoring_results_p:
        log(INFO_MESSAGES['LOCATING_TITLE'].format(monitoring_title), 'info')
        cursor_element = monitoring_results_p._p
        
        for bridge_name in sorted(all_trend_results.keys()):
            log(INFO_MESSAGES['GENERATING_CHART'].format(bridge_name), 'info')
            sensors_data = all_trend_results[bridge_name]
            cursor_element = generate_bridge_content(
                doc, bridge_name, sensors_data, all_bridges_data, 
                cursor_element, pic_style, report_type
            )
    else:
        log(ERROR_MESSAGES['MONITORING_TITLE_NOT_FOUND'].format(monitoring_title), 'warn')
    
    # 生成结论
    if conclusion_p:
        log(INFO_MESSAGES['LOCATING_TITLE'].format('结论'), 'info')
        
        cursor_element = conclusion_p._p
        generate_conclusion(doc, all_trend_results, all_bridges_data, cursor_element)
    else:
        log(ERROR_MESSAGES['CONCLUSION_TITLE_NOT_FOUND'], 'warn')
    
    # 保存文档
    try:
        doc.save(report_path)
        report_name = "月报" if report_type == 'MONTH' else "季报"
        log(SUCCESS_MESSAGES['REPORT_SAVED'].format(report_name, report_path))
        return report_path
    except Exception as e:
        report_name = "月报" if report_type == 'MONTH' else "季报"
        log(ERROR_MESSAGES['SAVE_FAILED'].format(report_name, str(e)))
        return None

def generate_month_report(all_trend_results, all_bridges_data, output_dir):
    """生成月报"""
    return generate_report_base(all_trend_results, all_bridges_data, output_dir, 'MONTH')

def generate_quarter_report(all_trend_results, all_bridges_data, output_dir):
    """生成季报"""
    return generate_report_base(all_trend_results, all_bridges_data, output_dir, 'QUARTER')

def generate_all_trend_charts_report(all_trend_results, all_bridges_data, output_dir, manufacturers_map=None):
    """
    生成所有桥梁所有测点趋势图一览表Word报告。
    每个桥梁分区块，区块内为该桥所有测点的趋势图表格：
    - 每行6个图
    - 下一行显示图片下方的测点名
    - 再下一行显示厂家信息（若提供manufacturers_map）
    """
    if not DOCX_AVAILABLE:
        log(ERROR_MESSAGES['DOCX_NOT_AVAILABLE'])
        return None

    # 导入gc用于内存管理
    import gc
    
    log("开始生成所有桥梁测点趋势图一览...", 'info')
    doc = Document()
    doc.add_heading('所有桥梁测点趋势图一览', level=1)
    doc.add_paragraph()  # 标题下空行
    pic_per_row = 6
    img_height = Cm(1.8)
    temp_imgs = []
    
    # 统计总测点数
    total_sensors_all = sum(len(sensors) for sensors in all_trend_results.values())
    processed_sensors_all = 0

    for bridge_idx, bridge_name in enumerate(sorted(all_trend_results.keys()), 1):
        log(f"正在处理桥梁 ({bridge_idx}/{len(all_trend_results)}): {bridge_name}", 'info')
        doc.add_heading(bridge_name, level=2)
        doc.add_paragraph()  # 桥标题下空行
        sensors_data = all_trend_results[bridge_name]
        bridge_df_dict = all_bridges_data[bridge_name]
        sensor_ids = sorted(sensors_data.keys())
        n = len(sensor_ids)
        if n == 0:
            doc.add_paragraph('无测点数据')
            continue
        # 计算表格行数：每6个图一行 + 一行测点名 + 一行厂家
        rows = ((n + pic_per_row - 1) // pic_per_row) * 3
        table = doc.add_table(rows=rows, cols=pic_per_row)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        # 填充图片、测点名与厂家
        # 准备临时目录
        temp_dir = os.path.join(output_dir, "临时图片")
        os.makedirs(temp_dir, exist_ok=True)
        total_sensors = len(sensor_ids)
        for idx, sensor_id in enumerate(sensor_ids):
            # 每处理10个测点输出一次进度
            if (idx + 1) % 10 == 0 or (idx + 1) == total_sensors:
                log(f"正在处理 {bridge_name} 的测点: {idx + 1}/{total_sensors}", 'info')
            
            row_img = (idx // pic_per_row) * 3
            col = idx % pic_per_row
            df = bridge_df_dict.get(sensor_id)
            trend_result = sensors_data.get(sensor_id)
            temp_img_path = None
            # 图片单元格
            cell = table.cell(row_img, col)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            # 判断是否离线/无效
            is_offline = (trend_result == "OFFLINE") or (not isinstance(trend_result, dict)) or (df is None)
            if not is_offline:
                try:
                    temp_img_path = os.path.join(temp_dir, f"temp_{bridge_name}_{sensor_id}_trend.png")
                    create_trend_chart(df, trend_result, temp_img_path)
                    temp_imgs.append(temp_img_path)
                    run = p.add_run()
                    run.add_picture(temp_img_path, height=img_height)
                    # 立即清理matplotlib缓存，释放内存
                    import matplotlib.pyplot as plt
                    plt.clf()  # 清除当前figure
                    gc.collect()  # 强制垃圾回收
                    # 立即清理matplotlib缓存，释放内存
                    import matplotlib.pyplot as plt
                    plt.clf()  # 清除当前figure
                    import gc
                    gc.collect()  # 强制垃圾回收
                except Exception as e:
                    log(f"生成测点 {sensor_id} 的图表失败: {e}", 'warn')
                    # 失败时显示占位符
                    placeholder = p.add_run("生成失败")
                    try:
                        placeholder.font.color.rgb = RGBColor(255, 0, 0)
                        placeholder.bold = True
                    except Exception:
                        pass
            else:
                # 离线占位
                placeholder = p.add_run("离线")
                try:
                    placeholder.font.color.rgb = RGBColor(128, 128, 128)
                    placeholder.bold = True
                except Exception:
                    pass
            # 下方一行写测点名
            cell_title = table.cell(row_img + 1, col)
            p_title = cell_title.paragraphs[0]
            p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_title.add_run(sensor_id)
            # 再下一行写厂家
            manufacturer = ''
            if manufacturers_map and bridge_name in manufacturers_map:
                manufacturer = manufacturers_map[bridge_name].get(sensor_id, '')
            cell_vendor = table.cell(row_img + 2, col)
            p_vendor = cell_vendor.paragraphs[0]
            p_vendor.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_vendor.add_run(manufacturer if manufacturer else ("离线" if is_offline else ""))
        add_table_border(table)
        doc.add_paragraph()  # 桥之间空行
        
        # 每处理完一个桥梁，清理一次内存
        processed_sensors_all += n
        log(f"已完成 {bridge_name} 的 {n} 个测点，总进度: {processed_sensors_all}/{total_sensors_all}", 'info')
        gc.collect()  # 强制垃圾回收
    
    # 保存文档
    report_path = os.path.join(output_dir, '所有桥梁测点趋势图一览.docx')
    try:
        doc.save(report_path)
        log(f"所有桥梁测点趋势图一览已保存至: {report_path}", 'success')
    except Exception as e:
        log(f"保存所有测点趋势图一览失败: {e}", 'error')
        report_path = None
    # 清理临时图片
    for img in temp_imgs:
        try:
            if os.path.exists(img):
                os.remove(img)
        except Exception:
            pass
    return report_path

