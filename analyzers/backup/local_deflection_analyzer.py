import os
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import matplotlib.dates as mdates
from datetime import datetime, timedelta
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import warnings
warnings.filterwarnings('ignore')

plt.rcParams['font.sans-serif'] = ['SimSun']
plt.rcParams['axes.unicode_minus'] = False

class LocalDeflectionAnalyzer:
    def __init__(self):
        # 配置路径
        self.deflection_file = r"D:\03-chengxu\cursor\report_gen\数据下载\白土北江特大桥年数据\挠度\数据\2025-07-31-09-16-08左幅挠度测点PRE-L3_.txt"
        self.output_dir = r"D:\03-chengxu\cursor\report_gen\naodu"
        self.bridge_name = "白土北江特大桥"
        self.deflection_sensor_name = "PRE-L6"
        self.template_path = r"D:\03-chengxu\cursor\report_gen\年报模版.docx"
        
        # 局部分析配置
        self.time_window_minutes = 5  # 前后半分钟，可轻松更改
        
        # 数据存储
        self.deflection_data = None
        self.selected_time = None
        self.local_data = None
        
        # 创建输出目录
        if not os.path.exists(self.output_dir):
            os.makedirs(self.output_dir)

    def load_data(self):
        print("正在加载数据...")
        try:
            df = pd.read_csv(self.deflection_file, sep='\s+', header=None, names=['Date', 'Time', 'Deflection'], encoding='utf-8')
            print("挠度前5行：", df.head())
            df['Time'] = pd.to_datetime(df['Date'] + ' ' + df['Time'], format='%Y-%m-%d %H:%M:%S.%f')
            self.deflection_data = df[['Time', 'Deflection']].sort_values('Time')
            print(f"✅ 成功加载挠度文件: {os.path.basename(self.deflection_file)} 共{len(self.deflection_data)}条")
        except Exception as e:
            print(f"❌ 加载挠度文件失败: {e}")
            return False
        
        return True

    def select_time_for_analysis(self):
        print(f"数据时间范围: {self.deflection_data['Time'].min()} 到 {self.deflection_data['Time'].max()}")
        
        # 获取可用日期
        available_dates = self.deflection_data['Time'].dt.date.unique()
        available_dates = sorted(available_dates)
        
        print(f"📅 可用日期范围: {available_dates[0]} 到 {available_dates[-1]}")
        print(f"💡 提示：请输入时间格式如 2025-07-15 08:30:00")
        
        # 设置默认时间
        default_date = available_dates[len(available_dates)//2]
        default_time = f"{default_date} 12:00:00"
        
        while True:
            try:
                time_input = input(f"\n🔍 请输入要分析的时间点 (默认: {default_time}): ").strip()
                
                if not time_input:  # 如果用户直接按回车
                    print(f"🔄 使用默认时间: {default_time}")
                    self.selected_time = datetime.strptime(default_time, '%Y-%m-%d %H:%M:%S')
                    return True
                
                try:
                    selected_time = datetime.strptime(time_input, '%Y-%m-%d %H:%M:%S')
                    
                    # 检查时间是否在数据范围内
                    min_time = self.deflection_data['Time'].min()
                    max_time = self.deflection_data['Time'].max()
                    
                    if min_time <= selected_time <= max_time:
                        self.selected_time = selected_time
                        print(f"✅ 已选择时间: {self.selected_time}")
                        return True
                    else:
                        print(f"❌ 时间 {selected_time} 不在可用范围内")
                        print(f"💡 可用时间范围: {min_time} 到 {max_time}")
                except ValueError:
                    print("❌ 时间格式错误，请使用 YYYY-MM-DD HH:MM:SS 格式")
                    print("💡 示例: 2025-07-15 08:30:00")
            except (EOFError, KeyboardInterrupt):
                print(f"\n🔄 自动选择默认时间: {default_time}")
                self.selected_time = datetime.strptime(default_time, '%Y-%m-%d %H:%M:%S')
                print(f"✅ 已选择时间: {self.selected_time}")
                return True
            except Exception as e:
                print(f"❌ 输入错误: {e}")
                print(f"🔄 使用默认时间: {default_time}")
                self.selected_time = datetime.strptime(default_time, '%Y-%m-%d %H:%M:%S')
                return True

    def extract_local_data(self):
        print(f"正在提取 {self.selected_time} 前后{self.time_window_minutes}分钟的数据...")
        
        # 计算时间窗口
        start_time = self.selected_time - timedelta(minutes=self.time_window_minutes)
        end_time = self.selected_time + timedelta(minutes=self.time_window_minutes)
        
        print(f"提取时间范围: {start_time} 到 {end_time}")
        
        # 检查数据采样频率
        sample_interval = self.deflection_data['Time'].diff().median()
        print(f"数据平均采样间隔: {sample_interval}")
        
        # 估算应该有多少个数据点
        expected_points = (self.time_window_minutes * 2 * 60) / sample_interval.total_seconds()
        print(f"预期数据点数: {expected_points:.0f}")
        
        local_deflection = self.deflection_data[
            (self.deflection_data['Time'] >= start_time) & 
            (self.deflection_data['Time'] <= end_time)
        ].copy()
        
        if len(local_deflection) == 0:
            print(f"❌ 未找到 {self.selected_time} 附近的挠度数据")
            return False
        
        # 显示提取数据的信息
        if len(local_deflection) > 0:
            min_time = local_deflection['Time'].min()
            max_time = local_deflection['Time'].max()
            time_span = max_time - min_time
            print(f"提取数据时间范围: {min_time} 到 {max_time}")
            print(f"数据时间跨度: {time_span}")
            print(f"实际数据点数: {len(local_deflection)}")
            
            # 如果数据点太少，尝试扩大时间窗口
            if len(local_deflection) < 50:  # 如果少于50个点
                print(f"⚠️ 数据点较少，尝试扩大时间窗口...")
                expanded_start = self.selected_time - timedelta(minutes=self.time_window_minutes * 2)
                expanded_end = self.selected_time + timedelta(minutes=self.time_window_minutes * 2)
                
                expanded_deflection = self.deflection_data[
                    (self.deflection_data['Time'] >= expanded_start) & 
                    (self.deflection_data['Time'] <= expanded_end)
                ].copy()
                
                if len(expanded_deflection) > len(local_deflection):
                    print(f"✅ 扩大时间窗口后获得 {len(expanded_deflection)} 个数据点")
                    local_deflection = expanded_deflection
                    self.time_window_minutes *= 2  # 更新显示的时间窗口
        
        self.local_data = local_deflection
        
        print(f"✅ 成功提取局部数据: {len(local_deflection)}条")
        return True

    def process_local_data(self):
        """处理局部数据，参考MATLAB代码"""
        print("正在处理局部数据...")
        
        # 参考MATLAB代码处理
        deflection_values = self.local_data['Deflection'].values
        time_values = self.local_data['Time'].values
        
        # MATLAB: dis=-cell2mat(table2cell(T(:,2)))
        deflection_values = -deflection_values
        
        # MATLAB: dis=dis-dis(1)
        deflection_values = deflection_values - deflection_values[0]
        
        print(f"处理后数据范围: {np.min(deflection_values):.2f} 到 {np.max(deflection_values):.2f}")
        
        return {
            'time': time_values,
            'deflection': deflection_values
        }

    def create_local_plot(self, data):
        print("正在生成局部放大图...")
        
        fig, ax = plt.subplots(1, 1, figsize=(12, 8))
        
        # 获取时间字符串
        time_str = self.selected_time.strftime('%Y年%m月%d日 %H:%M:%S')
        
        # 绘制局部数据
        ax.plot(data['time'], data['deflection'], 'b', linewidth=1)
        ax.set_xlabel('时间', fontsize=15)
        ax.set_ylabel('挠度（mm）', fontsize=15)
        ax.set_title(f'{self.bridge_name} {self.deflection_sensor_name} 挠度监测瞬时数据', fontsize=15)
        ax.grid(True, linestyle=':', alpha=0.5)
        ax.tick_params(axis='both', which='major', labelsize=15)
        
        # 在图上标记选中的时间点
        ax.axvline(x=self.selected_time, color='r', linestyle='--', alpha=0.7, label=f'选中时间点: {self.selected_time.strftime("%H:%M:%S")}')
        ax.legend()
        
        # 添加时间窗口信息
        ax.text(0.02, 0.98, f"时间窗口: ±{self.time_window_minutes}分钟", 
                transform=ax.transAxes, fontsize=10, verticalalignment='top', 
                bbox=dict(boxstyle='round', facecolor='white', alpha=0.8))
        
        # 设置x轴格式
        ax.xaxis.set_major_formatter(mdates.DateFormatter('%H:%M:%S'))
        ax.xaxis.set_major_locator(mdates.SecondLocator(interval=30))
        plt.setp(ax.xaxis.get_majorticklabels(), rotation=45)
        
        plt.tight_layout()
        
        # 使用非阻塞模式显示图片（不保存）
        plt.show(block=False)
        plt.pause(0.1)  # 短暂暂停以确保图片显示

    def generate_local_report(self, data, time_str):
        print("正在生成局部分析报告...")
        if not os.path.exists(self.template_path):
            print(f"❌ 模板文件不存在: {self.template_path}")
            return None
        
        doc = Document(self.template_path)
        
        # 添加标题
        title = doc.add_heading(f'{self.bridge_name} {self.deflection_sensor_name} 局部挠度分析报告', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 添加分析信息
        doc.add_heading('分析信息', level=1)
        self.add_field(doc, "桥梁名称", self.bridge_name)
        self.add_field(doc, "测点编号", self.deflection_sensor_name)
        self.add_field(doc, "分析时间点", time_str)
        self.add_field(doc, "时间窗口", f"±{self.time_window_minutes}分钟")
        self.add_field(doc, "数据点数", len(self.local_data))
        
        # 添加统计结果
        deflection_values = data['deflection']
        doc.add_heading('统计结果', level=1)
        self.add_field(doc, "最大挠度", f"{np.max(deflection_values):.2f}", " mm")
        self.add_field(doc, "最小挠度", f"{np.min(deflection_values):.2f}", " mm")
        self.add_field(doc, "挠度变化范围", f"{np.max(deflection_values) - np.min(deflection_values):.2f}", " mm")
        self.add_field(doc, "平均挠度", f"{np.mean(deflection_values):.2f}", " mm")
        
        # 添加图表
        doc.add_heading('局部放大图', level=1)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        image_path = os.path.join(self.output_dir, f"{time_str.replace(':', '-')}_局部挠度分析.png")
        p.add_run().add_picture(image_path, width=Inches(6))
        self.add_caption(doc, '图', f'{time_str} 局部挠度分析')
        
        # 保存报告
        report_path = os.path.join(self.output_dir, f'{self.bridge_name}_{self.deflection_sensor_name}_{time_str.replace(":", "-")}_局部挠度分析报告.docx')
        doc.save(report_path)
        print(f"✅ 报告已保存: {report_path}")
        return report_path

    def add_caption(self, doc, caption_type, caption_text):
        """添加图片标题"""
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"{caption_type} {caption_text}")
        run.font.size = doc.styles['Normal'].font.size
        run.font.bold = True

    def add_field(self, doc, field_name, field_value, unit=""):
        """添加字段到文档"""
        p = doc.add_paragraph()
        p.add_run(f"{field_name}: ").bold = True
        p.add_run(f"{field_value}{unit}")

    def run_interactive_analysis(self):
        print("开始局部挠度分析...")
        
        if not self.load_data():
            return False
        
        while True:
            if not self.select_time_for_analysis():
                continue
            
            if not self.extract_local_data():
                continue
            
            # 处理局部数据
            processed_data = self.process_local_data()
            
            # 生成局部图表（不保存）
            self.create_local_plot(processed_data)
            
            print(f"\n📊 {self.selected_time} 局部分析结果:")
            deflection_values = processed_data['deflection']
            print(f"  最大挠度: {np.max(deflection_values):.2f} mm")
            print(f"  最小挠度: {np.min(deflection_values):.2f} mm")
            print(f"  挠度变化范围: {np.max(deflection_values) - np.min(deflection_values):.2f} mm")
            print(f"  平均挠度: {np.mean(deflection_values):.2f} mm")
            
            # 询问是否满意当前结果
            try:
                satisfaction_choice = input("\n💭 您对当前分析结果满意吗？(y/n): ")
                if satisfaction_choice.lower() in ['y', 'yes', '是', '满意']:
                    print("✅ 满意！正在保存结果并生成报告...")
                    
                    # 保存图片
                    time_str = self.selected_time.strftime('%Y-%m-%d %H-%M-%S')
                    output_path = os.path.join(self.output_dir, f"{time_str}_局部挠度分析.png")
                    plt.savefig(output_path, dpi=300, bbox_inches='tight')
                    print(f"✅ 局部分析图表已保存: {output_path}")
                    
                    # 生成报告
                    report_path = self.generate_local_report(processed_data, time_str)
                    print(f"✅ 报告已生成: {report_path}")
                    
                    # 询问是否继续分析其他时间点
                    try:
                        continue_choice = input("\n是否继续分析其他时间点？(y/n): ")
                        if continue_choice.lower() in ['y', 'yes', '是']:
                            continue
                        else:
                            print("👋 分析结束")
                            return True
                    except (EOFError, KeyboardInterrupt):
                        print("\n👋 分析结束")
                        return True
                else:
                    print("🔄 不满意，请选择其他时间点重新分析...")
                    # 关闭当前图片窗口
                    plt.close('all')
                    continue
            except (EOFError, KeyboardInterrupt):
                print("\n👋 分析结束")
                return True

def main():
    print("=" * 50)
    print("局部挠度分析器启动中...")
    print("=" * 50)
    
    try:
        analyzer = LocalDeflectionAnalyzer()
        print("✅ 分析器初始化成功")
        result = analyzer.run_interactive_analysis()
        print(f"✅ 分析完成，结果: {result}")
    except Exception as e:
        print(f"❌ 程序运行出错: {e}")
        import traceback
        traceback.print_exc()

if __name__ == "__main__":
    main() 