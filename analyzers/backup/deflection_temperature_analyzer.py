import os
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import matplotlib.dates as mdates
from datetime import datetime, timedelta
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import warnings
warnings.filterwarnings('ignore')

# 使用滑动平均方法进行温度趋势提取
print("使用滑动平均方法进行挠度温度效应分离")

plt.rcParams['font.sans-serif'] = ['SimSun']
plt.rcParams['axes.unicode_minus'] = False

class DeflectionTemperatureAnalyzer:
    def __init__(self):
        # 配置路径
        # self.deflection_file = r"D:\03-chengxu\cursor\report_gen\数据下载\白土北江特大桥年数据\挠度\数据\20250730_160639_654_白土北江特大桥_654_PRE-L6_左幅挠度_挠度_左幅_主跨跨中_单日_温度分离_瞬时.txt"
        self.deflection_file = r"D:\03-chengxu\cursor\report_gen\数据下载\白土北江特大桥\挠度\数据\2023-04-08-12-13-15交科悬索桥连通管测点PRE-L4（跨中下游）_.txt"
        
        
        self.output_dir = r"D:\03-chengxu\cursor\report_gen\naodu"
        self.bridge_name = "白土北江特大桥"
        self.deflection_sensor_name = "PRE-L6"
        self.template_path = r"D:\03-chengxu\cursor\report_gen\年报模版.docx"
        
        # 滑动窗口大小配置（可调整）
        self.sliding_window_size =215  # 可以根据需要调整这个值
        
        # 数据存储
        self.deflection_data = None
        self.selected_date = None
        self.daily_deflection = None
        
        # 创建输出目录
        if not os.path.exists(self.output_dir):
            os.makedirs(self.output_dir)

    def load_data(self):
        print("正在加载数据...")
        # 加载挠度数据
        try:
            # 读取数据，处理不同的时间格式
            df = pd.read_csv(self.deflection_file, sep='\s+', header=None, names=['DateTime', 'Deflection'], encoding='utf-8')
            print("挠度前5行：", df.head())
            
            # 尝试不同的时间格式
            try:
                # 尝试格式1: YYYY-MM-DD HH:MM:SS.fff
                df['Time'] = pd.to_datetime(df['DateTime'], format='%Y-%m-%d %H:%M:%S.%f')
            except:
                try:
                    # 尝试格式2: YYYY/M/D H:M:S
                    df['Time'] = pd.to_datetime(df['DateTime'], format='%Y/%m/%d %H:%M:%S')
                except:
                    # 尝试格式3: 自动解析
                    df['Time'] = pd.to_datetime(df['DateTime'])
            
            self.deflection_data = df[['Time', 'Deflection']].sort_values('Time')
            print(f"✅ 成功加载挠度文件: {os.path.basename(self.deflection_file)} 共{len(self.deflection_data)}条")
        except Exception as e:
            print(f"❌ 加载挠度文件失败: {e}")
            return False
        
        return True

    def select_date_for_analysis(self):
        print(f"数据时间范围: {self.deflection_data['Time'].min()} 到 {self.deflection_data['Time'].max()}")
        
        # 获取可用日期
        available_dates = self.deflection_data['Time'].dt.date.unique()
        available_dates = sorted(available_dates)
        
        print(f"📅 可用日期范围: {available_dates[0]} 到 {available_dates[-1]}")
        print(f"📅 共有 {len(available_dates)} 天的数据")
        print(f"💡 提示：请输入日期格式如 2025-07-15，或输入 'list' 查看所有日期")
        
        # 设置默认日期
        default_date = available_dates[len(available_dates)//2]
        
        while True:
            try:
                date_input = input("\n🔍 请输入要分析的日期: ").strip()
                
                if not date_input:  # 如果用户直接按回车
                    print(f"🔄 使用默认日期: {default_date}")
                    self.selected_date = default_date
                    return True
                
                if date_input.lower() == 'list':
                    print("\n📋 可用日期列表:")
                    for i, date in enumerate(available_dates):
                        if i % 10 == 0 and i > 0:
                            print()
                        print(f"{date}", end=" ")
                    print("\n")
                    continue
                
                try:
                    selected_date = datetime.strptime(date_input, '%Y-%m-%d').date()
                    if selected_date in available_dates:
                        self.selected_date = selected_date
                        print(f"✅ 已选择日期: {self.selected_date}")
                        return True
                    else:
                        print(f"❌ 日期 {selected_date} 不在可用范围内")
                        print(f"💡 可用日期范围: {available_dates[0]} 到 {available_dates[-1]}")
                except ValueError:
                    print("❌ 日期格式错误，请使用 YYYY-MM-DD 格式")
                    print("💡 示例: 2025-07-15")
            except (EOFError, KeyboardInterrupt):
                # 当没有输入时，自动选择默认日期
                print(f"\n🔄 自动选择默认日期: {default_date}")
                self.selected_date = default_date
                print(f"✅ 已选择日期: {self.selected_date}")
                return True
            except Exception as e:
                print(f"❌ 输入错误: {e}")
                print(f"🔄 使用默认日期: {default_date}")
                self.selected_date = default_date
                return True

    def extract_daily_data(self):
        print(f"正在提取 {self.selected_date} 的数据...")
        
        # 提取指定日期的数据
        start_time = pd.Timestamp(self.selected_date)
        end_time = start_time + timedelta(days=1)
        
        print(f"提取时间范围: {start_time} 到 {end_time}")
        
        daily_deflection = self.deflection_data[
            (self.deflection_data['Time'] >= start_time) & 
            (self.deflection_data['Time'] < end_time)
        ].copy()
        
        if len(daily_deflection) == 0:
            print(f"❌ 未找到 {self.selected_date} 的挠度数据")
            return False
        
        # 显示提取数据的时间范围
        if len(daily_deflection) > 0:
            min_time = daily_deflection['Time'].min()
            max_time = daily_deflection['Time'].max()
            time_span = max_time - min_time
            print(f"提取数据时间范围: {min_time} 到 {max_time}")
            print(f"数据时间跨度: {time_span}")
            print(f"是否覆盖整天: {'是' if time_span >= timedelta(hours=23) else '否'}")
            print(f"数据点间隔: {(max_time - min_time) / len(daily_deflection)}")
        
        self.daily_deflection = daily_deflection
        
        print(f"✅ 成功提取 {self.selected_date} 的数据: 挠度{len(daily_deflection)}条")
        return True

    def smooth_data(self, data, window_size):
        """滑动平均平滑数据 - 改进的边界处理"""
        if window_size >= len(data):
            return data
        
        # 使用pandas的rolling方法，提供更好的边界处理
        smoothed = pd.Series(data).rolling(window=window_size, center=True, min_periods=1).mean().values
        
        # 如果pandas方法不可用，使用卷积方法
        if np.any(np.isnan(smoothed)):
            # 使用卷积进行滑动平均
            kernel = np.ones(window_size) / window_size
            smoothed = np.convolve(data, kernel, mode='same')
            
            # 改进的边界处理：使用渐变的窗口大小
            half_window = window_size // 2
            for i in range(half_window):
                # 左边界：使用逐渐增大的窗口
                left_window = min(i + 1, half_window)
                smoothed[i] = np.mean(data[:left_window])
                
                # 右边界：使用逐渐减小的窗口
                right_window = min(len(data) - i - 1, half_window)
                smoothed[-(i+1)] = np.mean(data[-right_window:])
        
        return smoothed

    def separate_deflection_components(self):
        print("正在进行挠度分量分离...")
        deflection_values = self.daily_deflection['Deflection'].values
        time_values = self.daily_deflection['Time'].values
        
        # 打印原始数据范围
        print(f"原始数据范围: {np.min(deflection_values):.2f} 到 {np.max(deflection_values):.2f}")
        
        # 保存原始数据用于显示
        original_deflection_raw = deflection_values.copy()
        
        # 直接从原始数据提取温度趋势
        print("使用滑动平均方法进行温度趋势提取...")
        window_size = self.sliding_window_size
        print(f"使用滑动窗口大小: {window_size}")
        temperature_trend = self.smooth_data(original_deflection_raw, window_size)
        
        print(f"温度趋势范围: {np.min(temperature_trend):.2f} 到 {np.max(temperature_trend):.2f}")
        
        # 车辆荷载挠度 = 原始数据 - 温度趋势
        vehicle_deflection = original_deflection_raw - temperature_trend
        
        print(f"车辆荷载挠度范围: {np.min(vehicle_deflection):.2f} 到 {np.max(vehicle_deflection):.2f}")
        
        return {
            'time': time_values,
            'original': original_deflection_raw,  # 使用原始数据
            'temperature_trend': temperature_trend,
            'vehicle_deflection': vehicle_deflection
        }

    def analyze_daily_statistics(self, components):
        """分析每日统计数据"""
        print("正在进行统计分析...")
        
        # 温度趋势统计
        temp_trend = components['temperature_trend']
        temp_min_idx = np.argmin(temp_trend)
        temp_max_idx = np.argmax(temp_trend)
        
        # 车辆荷载统计
        vehicle_deflection = components['vehicle_deflection']
        vehicle_min_idx = np.argmin(vehicle_deflection)
        vehicle_max_idx = np.argmax(vehicle_deflection)
        
        return {
            'temp_deflection_range': np.max(temp_trend) - np.min(temp_trend),
            'vehicle_deflection_range': np.max(vehicle_deflection) - np.min(vehicle_deflection),
            'temp_min_time': components['time'][temp_min_idx],
            'temp_max_time': components['time'][temp_max_idx],
            'temp_min_value': temp_trend[temp_min_idx],
            'temp_max_value': temp_trend[temp_max_idx],
            'vehicle_min_time': components['time'][vehicle_min_idx],
            'vehicle_max_time': components['time'][vehicle_max_idx],
            'vehicle_min_value': vehicle_deflection[vehicle_min_idx],
            'vehicle_max_value': vehicle_deflection[vehicle_max_idx]
        }

    def create_analysis_plots(self, components):
        print("正在生成分析图表...")
        fig, axes = plt.subplots(3, 1, figsize=(12, 10))
        
        # 获取日期字符串
        date_str = self.selected_date.strftime('%Y年%m月%d日')
        
        # 原始信号 - 严格按照MATLAB样式
        axes[0].plot(components['time'], components['original'], 'b', linewidth=1)
        axes[0].set_xlabel('时间', fontsize=15)
        axes[0].set_ylabel('实测总挠度（mm）', fontsize=15)
        axes[0].set_title('原始信号', fontsize=15)
        axes[0].grid(True, linestyle=':', alpha=0.5)
        axes[0].tick_params(axis='both', which='major', labelsize=15)
        axes[0].text(0.02, 0.98, date_str, 
                     transform=axes[0].transAxes, fontsize=10, verticalalignment='top', 
                     bbox=dict(boxstyle='round', facecolor='white', alpha=0.8))
        
        # 温度趋势 - 严格按照MATLAB样式
        axes[1].plot(components['time'], components['temperature_trend'], 'r', linewidth=1)
        axes[1].set_xlabel('时间', fontsize=15)
        axes[1].set_ylabel('温度挠度（mm）', fontsize=15)
        axes[1].set_title('温度趋势', fontsize=15)
        axes[1].grid(True, linestyle=':', alpha=0.5)
        axes[1].tick_params(axis='both', which='major', labelsize=15)
        
        # 车辆荷载挠度 - 严格按照MATLAB样式
        axes[2].plot(components['time'], components['vehicle_deflection'], 'k', linewidth=1)
        axes[2].set_xlabel('时间', fontsize=15)
        axes[2].set_ylabel('车辆荷载挠度（mm）', fontsize=15)
        axes[2].set_title('车辆荷载挠度', fontsize=15)
        axes[2].grid(True, linestyle=':', alpha=0.5)
        axes[2].tick_params(axis='both', which='major', labelsize=15)
        
        # 统一x轴和y轴格式
        for ax in axes:
            # x轴格式
            ax.xaxis.set_major_formatter(mdates.DateFormatter('%H:%M'))
            # 根据数据时间跨度调整刻度间隔
            time_span = pd.Timestamp(components['time'].max()) - pd.Timestamp(components['time'].min())
            if time_span <= pd.Timedelta(hours=2):
                ax.xaxis.set_major_locator(mdates.MinuteLocator(interval=30))
            elif time_span <= pd.Timedelta(hours=6):
                ax.xaxis.set_major_locator(mdates.HourLocator(interval=1))
            else:
                ax.xaxis.set_major_locator(mdates.HourLocator(interval=2))
            plt.setp(ax.xaxis.get_majorticklabels(), rotation=45)
            
            # y轴格式 - 统一为整数显示
            ax.yaxis.set_major_formatter(plt.FuncFormatter(lambda x, p: f'{int(x)}'))
        
        plt.tight_layout()
        
        # 使用非阻塞模式显示图片（不保存）
        plt.show(block=False)
        plt.pause(0.1)  # 短暂暂停以确保图片显示



    def add_caption(self, doc, caption_type, caption_text):
        """添加图片标题"""
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"{caption_type} {caption_text}")
        run.font.size = doc.styles['Normal'].font.size
        run.font.bold = True

    def add_field(self, doc, field_name, field_value, unit=""):
        """添加字段到文档"""
        p = doc.add_paragraph()
        p.add_run(f"{field_name}: ").bold = True
        p.add_run(f"{field_value}{unit}")

    def generate_dynamic_description(self, stats):
        """生成动态描述文字"""
        temp_min_time = pd.Timestamp(stats['temp_min_time']).strftime('%H:%M')
        temp_max_time = pd.Timestamp(stats['temp_max_time']).strftime('%H:%M')
        temp_range = stats['temp_deflection_range']
        vehicle_range = stats['vehicle_deflection_range']
        
        # 确定哪个效应更大
        if temp_range > vehicle_range:
            comparison = "温度效应引起的挠度变化大于车辆荷载引起的挠度变化"
        elif temp_range < vehicle_range:
            comparison = "车辆荷载引起的挠度变化大于温度效应引起的挠度变化"
        else:
            comparison = "温度效应和车辆荷载引起的挠度变化相当"
        
        description = f"由图可知，主跨跨中挠度在{temp_min_time}达到单日内最低点，{temp_max_time}达到单日内最高点。由温度效应引起的挠度变化在单天内约{temp_range:.1f}mm，而由车辆荷载引起的挠度峰值变化约{vehicle_range:.1f}mm，可见，对于该桥而言，单日内{comparison}。"
        
        return description

    def generate_analysis_report(self, stats, date_str):
        print("正在生成分析报告...")
        if not os.path.exists(self.template_path):
            print(f"❌ 模板文件不存在: {self.template_path}")
            return None
        
        doc = Document(self.template_path)
        
        # 添加标题
        title = doc.add_heading(f'{self.bridge_name} {self.deflection_sensor_name} 挠度温度效应分析报告', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 添加分析结果
        doc.add_heading('分析结果', level=1)
        self.add_field(doc, "温度效应挠度变化范围", f"{stats['temp_deflection_range']:.2f}", " mm")
        self.add_field(doc, "车辆荷载挠度变化范围", f"{stats['vehicle_deflection_range']:.2f}", " mm")
        
        # 添加动态生成的描述
        doc.add_heading('分析结论', level=1)
        description = self.generate_dynamic_description(stats)
        p = doc.add_paragraph(description)
        
        # 添加图表
        doc.add_heading('分析图表', level=1)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        image_path = os.path.join(self.output_dir, f"{date_str}_挠度温度效应分离分析.png")
        p.add_run().add_picture(image_path, width=Inches(6))
        self.add_caption(doc, '图', f'{self.selected_date} 挠度温度效应分离分析')
        
        # 保存报告
        report_path = os.path.join(self.output_dir, f'{self.bridge_name}_{self.deflection_sensor_name}_{self.selected_date}_挠度温度效应分析报告.docx')
        doc.save(report_path)
        print(f"✅ 报告已保存: {report_path}")
        return report_path

    def run_interactive_analysis(self):
        print("开始挠度温度效应分析...")
        
        if not self.load_data():
            return False
        
        while True:
            if not self.select_date_for_analysis():
                continue
            
            if not self.extract_daily_data():
                continue
            
            # 进行分量分离
            separated_data = self.separate_deflection_components()
            stats = self.analyze_daily_statistics(separated_data)
            
            # 生成图表（不保存）
            self.create_analysis_plots(separated_data)
            
            print(f"\n📊 {self.selected_date} 分析结果:")
            print(f"  温度效应挠度变化范围: {stats['temp_deflection_range']:.2f} mm")
            print(f"  车辆荷载挠度变化范围: {stats['vehicle_deflection_range']:.2f} mm")
            print(f"  温度效应最低点: {pd.Timestamp(stats['temp_min_time']).strftime('%H:%M')}")
            print(f"  温度效应最高点: {pd.Timestamp(stats['temp_max_time']).strftime('%H:%M')}")
            
            # 询问是否满意当前结果
            try:
                satisfaction_choice = input("\n💭 您对当前分析结果满意吗？(y/n): ")
                if satisfaction_choice.lower() in ['y', 'yes', '是', '满意']:
                    print("✅ 满意！正在保存结果并生成报告...")
                    
                    # 保存图片
                    output_path = os.path.join(self.output_dir, f"{self.selected_date.strftime('%Y-%m-%d')}_挠度温度效应分离分析.png")
                    plt.savefig(output_path, dpi=300, bbox_inches='tight')
                    print(f"✅ 分析图表已保存: {output_path}")
                    
                    # 生成报告
                    report_path = self.generate_analysis_report(stats, self.selected_date.strftime('%Y-%m-%d'))
                    print(f"✅ 报告已生成: {report_path}")
                    
                    # 询问是否继续分析其他日期
                    try:
                        continue_choice = input("\n是否继续分析其他日期？(y/n): ")
                        if continue_choice.lower() in ['y', 'yes', '是']:
                            continue
                        else:
                            print("👋 分析结束")
                            return True
                    except (EOFError, KeyboardInterrupt):
                        print("\n👋 分析结束")
                        return True
                else:
                    print("🔄 不满意，请选择其他日期重新分析...")
                    # 关闭当前图片窗口
                    plt.close('all')
                    continue
            except (EOFError, KeyboardInterrupt):
                print("\n👋 分析结束")
                return True

def main():
    print("=" * 50)
    print("挠度温度效应分析器启动中...")
    print("=" * 50)
    
    try:
        analyzer = DeflectionTemperatureAnalyzer()
        print("✅ 分析器初始化成功")
        result = analyzer.run_interactive_analysis()
        print(f"✅ 分析完成，结果: {result}")
    except Exception as e:
        print(f"❌ 程序运行出错: {e}")
        import traceback
        traceback.print_exc()

if __name__ == "__main__":
    main() 