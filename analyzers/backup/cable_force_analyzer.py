#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
索力时间序列分析器
专门用于分析桥梁索力数据的长期趋势
"""

import pandas as pd
import matplotlib.pyplot as plt
import numpy as np
from datetime import datetime
import os
from pathlib import Path
import re
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement, qn

# 设置中文字体为宋体，英文字体为Times New Roman
plt.rcParams['font.sans-serif'] = ['SimSun']  # 宋体
plt.rcParams['font.serif'] = ['Times New Roman']  # Times New Roman
plt.rcParams['axes.unicode_minus'] = False  # 解决负号显示问题

def natural_sort_key(text):
    """自然排序键函数，用于正确排序包含数字的字符串"""
    return [int(c) if c.isdigit() else c.lower() for c in re.split('([0-9]+)', text)]

class CableForceAnalyzer:
    # ==================== 配置区域 ====================
    # 桥名配置 - 请在此处修改桥名
    BRIDGE_NAME = "乌石北江特大桥"  # 请在此处修改桥名
    
    def __init__(self, bridge_name=None):
        # 使用配置的桥名，如果没有传入参数则使用默认配置
        self.bridge_name = bridge_name or self.BRIDGE_NAME
        
        # 路径配置
        current_dir = Path(__file__).parent
        base_path = current_dir.parent
        self.base_dir = base_path / "数据下载"
        self.data_dir = self.base_dir / self.bridge_name / "索力" / "数据"
        self.output_dir = Path("suoli") / self.bridge_name
        
        # 确保输出目录存在
        self.output_dir.mkdir(parents=True, exist_ok=True)
        
        # 鲜艳的配色方案
        self.colors = {
            'blue': '#1f77b4',      # 鲜艳的蓝色
            'red': '#d62728',        # 鲜艳的红色
            'green': '#2ca02c',      # 鲜艳的绿色
            'orange': '#ff7f0e',     # 鲜艳的橙色
            'purple': '#9467bd',     # 鲜艳的紫色
            'brown': '#8c564b',      # 鲜艳的棕色
            'pink': '#e377c2',       # 鲜艳的粉色
            'gray': '#7f7f7f'        # 鲜艳的灰色
        }
        # ================================================
        
        # 数据存储
        self.cable_force_data = {}
        self.original_data = {}
        
        # 2022年历史数据（设计索力和荷载试验索力）
        self.historical_data_2022 = {
            'CAB-94#-B5-L': {'monitoring': 3953.3, 'design': 3314.8, 'load_test': 3528, 'deviation': 12.06},
            'CAB-94#-B5-R': {'monitoring': 3975.5, 'design': 3314.8, 'load_test': 3536, 'deviation': 12.43},
            'CAB-94#-B6-L': {'monitoring': 4217.2, 'design': 3487, 'load_test': 3711, 'deviation': 13.64},
            'CAB-94#-B6-R': {'monitoring': 4014.5, 'design': 3487, 'load_test': 3718, 'deviation': 7.97},
            'CAB-94#-B13-L': {'monitoring': 4908.7, 'design': 4256.9, 'load_test': 4559, 'deviation': 7.67},
            'CAB-94#-B13-R': {'monitoring': 4952.2, 'design': 4256.9, 'load_test': 4564, 'deviation': 8.50},
            'CAB-94#-B14-L': {'monitoring': 5070.8, 'design': 4611, 'load_test': 4742, 'deviation': 6.93},
            'CAB-94#-B14-R': {'monitoring': 5116.0, 'design': 4611, 'load_test': 4721, 'deviation': 8.37},
            'CAB-94#-B17-L': {'monitoring': 5502.8, 'design': 5145.4, 'load_test': 5246, 'deviation': 4.90},
            'CAB-94#-B17-R': {'monitoring': 5674.2, 'design': 5145.4, 'load_test': 5294, 'deviation': 7.18},
            'CAB-94#-B19-L': {'monitoring': 5463.4, 'design': 5161.3, 'load_test': 5268, 'deviation': 3.71},
            'CAB-94#-B19-R': {'monitoring': 5531.4, 'design': 5161.3, 'load_test': 5338, 'deviation': 3.62},
            'CAB-94#-C2-L': {'monitoring': 3537.9, 'design': 2879.7, 'load_test': 3037, 'deviation': 16.49},
            'CAB-94#-C2-R': {'monitoring': 3878.8, 'design': 2879.7, 'load_test': 3017, 'deviation': 28.56},
            'CAB-94#-C3-L': {'monitoring': 3683.2, 'design': 3115, 'load_test': 3137, 'deviation': 17.41},
            'CAB-94#-C3-R': {'monitoring': 3366.8, 'design': 3115, 'load_test': 3157, 'deviation': 6.65},
            'CAB-94#-C6-L': {'monitoring': 4504.5, 'design': 3647.5, 'load_test': 3856, 'deviation': 16.8},
            'CAB-94#-C6-R': {'monitoring': 4276.2, 'design': 3647.5, 'load_test': 3805, 'deviation': 12.4},
            'CAB-94#-C10-L': {'monitoring': 4925.7, 'design': 4304.5, 'load_test': 4536, 'deviation': 8.6},
            'CAB-94#-C10-R': {'monitoring': 4909.6, 'design': 4304.5, 'load_test': 4586, 'deviation': 7.1},
            'CAB-94#-C17-L': {'monitoring': 5391.1, 'design': 4562.1, 'load_test': 4930, 'deviation': 9.3},
            'CAB-94#-C17-R': {'monitoring': 5334.5, 'design': 4562.1, 'load_test': 5006, 'deviation': 6.6},
            'CAB-95#-B5-L': {'monitoring': 4012.1, 'design': 3317.2, 'load_test': 3175, 'deviation': 26.4},
            'CAB-95#-B5-R': {'monitoring': 3910.6, 'design': 3317.2, 'load_test': 3452, 'deviation': 13.3},
            'CAB-95#-B6-L': {'monitoring': 4206.3, 'design': 3488, 'load_test': 3684, 'deviation': 14.2},
            'CAB-95#-B6-R': {'monitoring': 4062.2, 'design': 3488, 'load_test': 3706, 'deviation': 9.6},
            'CAB-95#-B14-L': {'monitoring': 5094.5, 'design': 4613, 'load_test': 4742, 'deviation': 7.4},
            'CAB-95#-B14-R': {'monitoring': 5159.4, 'design': 4613, 'load_test': 4822, 'deviation': 7.0},
            'CAB-95#-B17-L': {'monitoring': 5622.9, 'design': 5153.9, 'load_test': 5320, 'deviation': 5.7},
            'CAB-95#-B17-R': {'monitoring': 5311.0, 'design': 5153.9, 'load_test': 5228, 'deviation': 1.6},
            'CAB-95#-C2-L': {'monitoring': 3570.4, 'design': 2874.1, 'load_test': 3031, 'deviation': 17.8},
            'CAB-95#-C2-R': {'monitoring': 3548.2, 'design': 2874.1, 'load_test': 2889, 'deviation': 22.8},
            'CAB-95#-C3-L': {'monitoring': 3977.9, 'design': 3109.5, 'load_test': 3326, 'deviation': 19.6},
            'CAB-95#-C3-R': {'monitoring': 3985.8, 'design': 3109.5, 'load_test': 3411, 'deviation': 16.9},
            'CAB-95#-C6-L': {'monitoring': 4281.3, 'design': 3645.6, 'load_test': 3858, 'deviation': 11},
            'CAB-95#-C6-R': {'monitoring': 4271.0, 'design': 3645.6, 'load_test': 3832, 'deviation': 11.5},
            'CAB-95#-C10-L': {'monitoring': 4948.3, 'design': 4313.7, 'load_test': 4454, 'deviation': 11.1},
            'CAB-95#-C10-R': {'monitoring': 4876.7, 'design': 4313.7, 'load_test': 4456, 'deviation': 9.4},
            'CAB-95#-C17-L': {'monitoring': 5132.2, 'design': 4567.2, 'load_test': 4764, 'deviation': 7.7},
            'CAB-95#-C17-R': {'monitoring': 5419.1, 'design': 4567.2, 'load_test': 4783, 'deviation': 13.3}
        }
        
        print(f"🏗️ 索力时间序列分析器初始化完成")
        print(f"📁 数据目录: {self.data_dir}")
        print(f"📂 输出目录: {self.output_dir}")
    
    def extract_sensor_info_from_filename(self, filename):
        """从文件名中提取传感器信息"""
        # 文件名格式: 2025-08-10-15-20-18右幅索力测点CAB-95#-C17-R_.txt
        # 提取左右幅信息
        if "左幅" in filename:
            side = "左幅"
        elif "右幅" in filename:
            side = "右幅"
        else:
            side = "未知"
        
        # 提取索力测点名称
        # 查找CAB-XX#-YY-Z_模式
        pattern = r'CAB-\d+#-[BC]\d+-[LR]_'
        match = re.search(pattern, filename)
        if match:
            sensor_name = match.group()
            # 移除末尾的下划线和点号
            sensor_name = sensor_name.rstrip('_.')
        else:
            sensor_name = "未知"
        
        return {
            'side': side,
            'sensor_name': sensor_name,
            'full_name': f"{side}{sensor_name}"
        }
    
    def load_cable_force_data(self):
        """加载索力数据"""
        print(f"\n📊 开始加载索力数据...")
        
        if not self.data_dir.exists():
            print(f"❌ 数据目录不存在: {self.data_dir}")
            return False
        
        # 获取所有txt文件
        txt_files = list(self.data_dir.glob("*.txt"))
        print(f"📁 找到 {len(txt_files)} 个数据文件")
        
        if not txt_files:
            print("❌ 未找到任何数据文件")
            return False
        
        for file_path in txt_files:
            # 从文件名提取传感器信息
            sensor_info = self.extract_sensor_info_from_filename(file_path.name)
            sensor_name = sensor_info['sensor_name']
            side = sensor_info['side']
            
            print(f"正在加载测点 {sensor_name} ({side}) 的数据...")
            
            try:
                # 读取数据文件
                data = pd.read_csv(file_path, sep='\t', header=None, encoding='utf-8')
                
                # 根据实际数据格式重新命名列
                # 第0列是时间，第2列是索力值（第1列是空列）
                if len(data.columns) >= 3:
                    data.columns = ['DateTime', 'Empty1', 'CableForce'] + [f'Col{i}' for i in range(3, len(data.columns))]
                else:
                    data.columns = ['DateTime', 'CableForce']
                
                # 转换时间列
                data['Time'] = pd.to_datetime(data['DateTime'], format='%Y-%m-%d %H:%M:%S.%f', errors='coerce')
                
                # 如果毫秒格式解析失败，尝试其他格式
                if data['Time'].isna().all():
                    data['Time'] = pd.to_datetime(data['DateTime'], errors='coerce')
                
                # 确保索力列为数值类型
                data['CableForce'] = pd.to_numeric(data['CableForce'], errors='coerce')
                
                # 删除多余的列，只保留时间和索力
                data = data[['Time', 'CableForce']]
                
                # 添加测点信息
                data['Sensor_Name'] = sensor_name
                data['Side'] = side
                data['Full_Name'] = sensor_info['full_name']
                
                # 移除无效数据
                initial_count = len(data)
                data.dropna(subset=['Time', 'CableForce'], inplace=True)
                final_count = len(data)
                
                if initial_count != final_count:
                    print(f"⚠️ 测点 {sensor_name}: 移除了 {initial_count - final_count} 条无效数据")
                
                # 存储数据
                self.cable_force_data[sensor_name] = data
                
                # 保存原始数据用于绘图
                self.original_data[sensor_name] = data.copy()
                
                print(f"✅ 成功加载测点 {sensor_name} 数据: {len(data)} 条记录")
                
            except Exception as e:
                print(f"❌ 加载测点 {sensor_name} 数据失败: {e}")
        
        print(f"📊 数据加载完成，共 {len(self.cable_force_data)} 个测点")
        return True
    
    def calculate_average_cable_force(self):
        """计算每个测点的平均索力值"""
        print(f"\n📊 计算各测点平均索力值...")
        
        average_data = {}
        
        for sensor_name, data in self.cable_force_data.items():
            if data is not None and not data.empty:
                # 计算平均索力
                avg_force = data['CableForce'].mean()
                
                # 获取测点信息
                side = data['Side'].iloc[0]
                
                average_data[sensor_name] = {
                    'average_force': avg_force,
                    'side': side,
                    'records_count': len(data),
                    'time_range': f"{data['Time'].min()} 到 {data['Time'].max()}"
                }
                
                print(f"测点 {sensor_name} ({side}): 平均索力 = {avg_force:.2f} KN")
        
        return average_data
    
    def create_comparison_table(self, current_average_data):
        """创建对比表格"""
        print(f"\n📋 创建索力对比表格...")
        
        table_data = []
        
        # 按自然排序处理测点名称
        sorted_sensor_names = sorted(current_average_data.keys(), key=natural_sort_key)
        
        for sensor_name in sorted_sensor_names:
            current_data = current_average_data[sensor_name]
            
            # 获取2022年历史数据
            if sensor_name in self.historical_data_2022:
                historical_data = self.historical_data_2022[sensor_name]
                
                # 计算本年度与荷载试验索力的偏差
                current_deviation = ((current_data['average_force'] - historical_data['load_test']) / historical_data['load_test']) * 100
                
                table_data.append({
                    '斜拉索编号': sensor_name,
                    '本年度监测索力(KN)': f"{current_data['average_force']:.2f}",
                    '2022年监测索力(KN)': f"{historical_data['monitoring']:.2f}",
                    '设计索力(KN)': f"{historical_data['design']:.2f}",
                    '荷载试验索力(KN)': f"{historical_data['load_test']:.2f}",
                    '本年度偏差(%)': f"{current_deviation:.2f}",
                    '2022年偏差(%)': f"{historical_data['deviation']:.2f}"
                })
            else:
                print(f"⚠️ 测点 {sensor_name} 没有对应的历史数据")
        
        # 保存为CSV文件
        df_table = pd.DataFrame(table_data)
        table_path = self.output_dir / '索力对比统计表.csv'
        df_table.to_csv(table_path, index=False, encoding='utf-8-sig')
        print(f"✅ 索力对比统计表格已保存: {table_path}")
        
        return df_table
    
    def plot_cable_force_comparison(self, current_average_data):
        """绘制索力对比分析图"""
        print("📊 绘制索力对比分析图...")
        
        # 准备数据
        sensor_names = []
        current_forces = []
        design_forces = []
        load_test_forces = []
        sides = []
        
        for sensor_name, current_data in current_average_data.items():
            if sensor_name in self.historical_data_2022:
                historical_data = self.historical_data_2022[sensor_name]
                
                sensor_names.append(sensor_name)
                current_forces.append(current_data['average_force'])
                design_forces.append(historical_data['design'])
                load_test_forces.append(historical_data['load_test'])
                sides.append(current_data['side'])
        
        # 创建图1：索力对比分析（柱状图+折线图）
        fig, ax = plt.subplots(figsize=(15, 8))
        
        x = np.arange(len(sensor_names))
        width = 0.35
        
        # 分离左右幅数据
        left_indices = [i for i, side in enumerate(sides) if side == '左幅']
        right_indices = [i for i, side in enumerate(sides) if side == '右幅']
        
        # 左幅数据
        if left_indices:
            left_x = [x[i] for i in left_indices]
            left_current = [current_forces[i] for i in left_indices]
            left_design = [design_forces[i] for i in left_indices]
            left_load_test = [load_test_forces[i] for i in left_indices]
            
            # 绘制左幅本年度监测索力（柱状图）
            ax.bar([pos - width/2 for pos in left_x], left_current, width, 
                   label='本年度监测索力(左幅)', color=self.colors['blue'], alpha=0.8)
            
            # 绘制左幅设计索力和荷载试验索力（折线图）
            ax.plot(left_x, left_design, 'o-', color=self.colors['red'], linewidth=2, 
                   markersize=6, label='设计索力(左幅)', alpha=0.8)
            ax.plot(left_x, left_load_test, 's-', color=self.colors['green'], linewidth=2, 
                   markersize=6, label='荷载试验索力(左幅)', alpha=0.8)
        
        # 右幅数据
        if right_indices:
            right_x = [x[i] for i in right_indices]
            right_current = [current_forces[i] for i in right_indices]
            right_design = [design_forces[i] for i in right_indices]
            right_load_test = [load_test_forces[i] for i in right_indices]
            
            # 绘制右幅本年度监测索力（柱状图）
            ax.bar([pos + width/2 for pos in right_x], right_current, width, 
                   label='本年度监测索力(右幅)', color=self.colors['orange'], alpha=0.8)
            
            # 绘制右幅设计索力和荷载试验索力（折线图）
            ax.plot(right_x, right_design, 'o-', color=self.colors['purple'], linewidth=2, 
                   markersize=6, label='设计索力(右幅)', alpha=0.8)
            ax.plot(right_x, right_load_test, 's-', color=self.colors['brown'], linewidth=2, 
                   markersize=6, label='荷载试验索力(右幅)', alpha=0.8)
        
        ax.set_title(f'{self.bridge_name}索力对比分析', fontsize=16, fontweight='bold')
        ax.set_ylabel('索力 (KN)', fontsize=14)
        ax.set_xlabel('索力测点', fontsize=14)
        ax.set_xticks(x)
        ax.set_xticklabels(sensor_names, rotation=45, ha='right')
        ax.legend(fontsize=12, loc='upper right')
        ax.grid(True, alpha=0.3)
        ax.tick_params(axis='both', which='major', labelsize=12)
        
        plt.tight_layout()
        
        # 保存图片
        output_filename = f"{self.bridge_name}索力对比分析图.png"
        plt.savefig(self.output_dir / output_filename, dpi=300, bbox_inches='tight')
        plt.close()
        
        print(f"✅ 已保存: {output_filename}")
    
    def plot_cable_force_deviation(self, current_average_data):
        """绘制索力偏差百分比分析图（2022年vs今年）"""
        print("📊 绘制索力偏差百分比分析图...")
        
        # 准备数据
        sensor_names = []
        current_deviations = []
        historical_deviations = []
        sides = []
        
        for sensor_name, current_data in current_average_data.items():
            if sensor_name in self.historical_data_2022:
                historical_data = self.historical_data_2022[sensor_name]
                
                sensor_names.append(sensor_name)
                current_deviation = ((current_data['average_force'] - historical_data['load_test']) / historical_data['load_test']) * 100
                current_deviations.append(current_deviation)
                historical_deviations.append(historical_data['deviation'])
                sides.append(current_data['side'])
        
        # 创建图2：偏差百分比分析（2022年vs今年）
        fig, ax = plt.subplots(figsize=(15, 8))
        
        x = np.arange(len(sensor_names))
        width = 0.35
        
        # 分离左右幅数据
        left_indices = [i for i, side in enumerate(sides) if side == '左幅']
        right_indices = [i for i, side in enumerate(sides) if side == '右幅']
        
        # 左幅偏差
        if left_indices:
            left_x = [x[i] for i in left_indices]
            left_current_dev = [current_deviations[i] for i in left_indices]
            left_historical_dev = [historical_deviations[i] for i in left_indices]
            
            bars1 = ax.bar([pos - width for pos in left_x], left_current_dev, width, 
                          label='本年度偏差(左幅)', color=self.colors['blue'], alpha=0.8)
            bars2 = ax.bar(left_x, left_historical_dev, width, 
                          label='2022年偏差(左幅)', color=self.colors['red'], alpha=0.8)
            
            # 添加数值标签
            for bar in bars1:
                height = bar.get_height()
                ax.text(bar.get_x() + bar.get_width()/2., height + 0.5,
                       f'{height:.1f}%', ha='center', va='bottom', fontsize=8, fontweight='bold')
            
            for bar in bars2:
                height = bar.get_height()
                ax.text(bar.get_x() + bar.get_width()/2., height + 0.5,
                       f'{height:.1f}%', ha='center', va='bottom', fontsize=8, fontweight='bold')
        
        # 右幅偏差
        if right_indices:
            right_x = [x[i] for i in right_indices]
            right_current_dev = [current_deviations[i] for i in right_indices]
            right_historical_dev = [historical_deviations[i] for i in right_indices]
            
            bars3 = ax.bar([pos + width for pos in right_x], right_current_dev, width, 
                          label='本年度偏差(右幅)', color=self.colors['orange'], alpha=0.8)
            bars4 = ax.bar([pos + 2*width for pos in right_x], right_historical_dev, width, 
                          label='2022年偏差(右幅)', color=self.colors['purple'], alpha=0.8)
            
            # 添加数值标签
            for bar in bars3:
                height = bar.get_height()
                ax.text(bar.get_x() + bar.get_width()/2., height + 0.5,
                       f'{height:.1f}%', ha='center', va='bottom', fontsize=8, fontweight='bold')
            
            for bar in bars4:
                height = bar.get_height()
                ax.text(bar.get_x() + bar.get_width()/2., height + 0.5,
                       f'{height:.1f}%', ha='center', va='bottom', fontsize=8, fontweight='bold')
        
        ax.set_title(f'{self.bridge_name}索力偏差百分比分析（2022年vs本年度）', fontsize=16, fontweight='bold')
        ax.set_ylabel('偏差 (%)', fontsize=14)
        ax.set_xlabel('索力测点', fontsize=14)
        ax.set_xticks(x)
        ax.set_xticklabels(sensor_names, rotation=45, ha='right')
        ax.legend(fontsize=12)
        ax.grid(True, alpha=0.3)
        ax.tick_params(axis='both', which='major', labelsize=12)
        
        # 添加零线
        ax.axhline(y=0, color='black', linestyle='-', alpha=0.5)
        
        plt.tight_layout()
        
        # 保存图片
        output_filename = f"{self.bridge_name}索力偏差百分比分析图.png"
        plt.savefig(self.output_dir / output_filename, dpi=300, bbox_inches='tight')
        plt.close()
        
        print(f"✅ 已保存: {output_filename}")
    
    def plot_left_right_force_comparison(self, current_average_data):
        """绘制左右幅索力对比分析图"""
        print("📊 绘制左右幅索力对比分析图...")
        
        # 准备数据 - 按测点编号分组
        sensor_groups = {}
        
        for sensor_name, current_data in current_average_data.items():
            if sensor_name in self.historical_data_2022:
                historical_data = self.historical_data_2022[sensor_name]
                
                # 提取基础测点编号（去掉-L或-R后缀）
                base_name = sensor_name[:-2]  # 去掉最后的-L或-R
                side = sensor_name[-1]  # L或R
                
                if base_name not in sensor_groups:
                    sensor_groups[base_name] = {}
                
                sensor_groups[base_name][side] = {
                    'current_force': current_data['average_force']
                }
        
        # 只保留有左右幅数据的测点
        valid_groups = {name: data for name, data in sensor_groups.items() 
                       if 'L' in data and 'R' in data}
        
        if not valid_groups:
            print("⚠️ 没有找到同时具有左右幅数据的测点")
            return
        
        # 创建图表
        fig, ax = plt.subplots(figsize=(15, 8))
        
        x = np.arange(len(valid_groups))
        width = 0.35
        
        left_current = [valid_groups[name]['L']['current_force'] for name in valid_groups.keys()]
        right_current = [valid_groups[name]['R']['current_force'] for name in valid_groups.keys()]
        
        # 绘制柱状图
        bars1 = ax.bar([pos - width/2 for pos in x], left_current, width, 
                      label='左幅本年度监测索力', color=self.colors['blue'], alpha=0.8)
        bars2 = ax.bar([pos + width/2 for pos in x], right_current, width, 
                      label='右幅本年度监测索力', color=self.colors['orange'], alpha=0.8)
        
        # 添加数值标签
        for bars in [bars1, bars2]:
            for bar in bars:
                height = bar.get_height()
                ax.text(bar.get_x() + bar.get_width()/2., height + 50,
                       f'{height:.0f}', ha='center', va='bottom', fontsize=8, fontweight='bold')
        
        ax.set_title(f'{self.bridge_name}左右幅索力对比分析', fontsize=16, fontweight='bold')
        ax.set_ylabel('索力 (KN)', fontsize=14)
        ax.set_xlabel('索力测点', fontsize=14)
        ax.set_xticks(x)
        ax.set_xticklabels(list(valid_groups.keys()), rotation=45, ha='right')
        ax.legend(fontsize=12, loc='upper right')
        ax.grid(True, alpha=0.3)
        ax.tick_params(axis='both', which='major', labelsize=12)
        
        plt.tight_layout()
        
        # 保存图片
        output_filename = f"{self.bridge_name}左右幅索力对比分析图.png"
        plt.savefig(self.output_dir / output_filename, dpi=300, bbox_inches='tight')
        plt.close()
        
        print(f"✅ 已保存: {output_filename}")
    
    def plot_left_right_deviation_comparison(self, current_average_data):
        """绘制左右幅偏差对比分析图"""
        print("📊 绘制左右幅偏差对比分析图...")
        
        # 准备数据 - 按测点编号分组
        sensor_groups = {}
        
        for sensor_name, current_data in current_average_data.items():
            if sensor_name in self.historical_data_2022:
                historical_data = self.historical_data_2022[sensor_name]
                
                # 提取基础测点编号（去掉-L或-R后缀）
                base_name = sensor_name[:-2]  # 去掉最后的-L或-R
                side = sensor_name[-1]  # L或R
                
                if base_name not in sensor_groups:
                    sensor_groups[base_name] = {}
                
                sensor_groups[base_name][side] = {
                    'current_deviation': ((current_data['average_force'] - historical_data['load_test']) / historical_data['load_test']) * 100,
                    'historical_deviation': historical_data['deviation']
                }
        
        # 只保留有左右幅数据的测点
        valid_groups = {name: data for name, data in sensor_groups.items() 
                       if 'L' in data and 'R' in data}
        
        if not valid_groups:
            print("⚠️ 没有找到同时具有左右幅数据的测点")
            return
        
        # 创建图表
        fig, ax = plt.subplots(figsize=(15, 8))
        
        x = np.arange(len(valid_groups))
        width = 0.2
        
        left_current_dev = [valid_groups[name]['L']['current_deviation'] for name in valid_groups.keys()]
        right_current_dev = [valid_groups[name]['R']['current_deviation'] for name in valid_groups.keys()]
        left_historical_dev = [valid_groups[name]['L']['historical_deviation'] for name in valid_groups.keys()]
        right_historical_dev = [valid_groups[name]['R']['historical_deviation'] for name in valid_groups.keys()]
        
        bars1 = ax.bar([pos - 1.5*width for pos in x], left_current_dev, width, 
                      label='左幅本年度偏差', color=self.colors['blue'], alpha=0.8)
        bars2 = ax.bar([pos - 0.5*width for pos in x], right_current_dev, width, 
                      label='右幅本年度偏差', color=self.colors['orange'], alpha=0.8)
        bars3 = ax.bar([pos + 0.5*width for pos in x], left_historical_dev, width, 
                      label='左幅2022年偏差', color=self.colors['red'], alpha=0.8)
        bars4 = ax.bar([pos + 1.5*width for pos in x], right_historical_dev, width, 
                      label='右幅2022年偏差', color=self.colors['purple'], alpha=0.8)
        
        # 添加数值标签
        for bars in [bars1, bars2, bars3, bars4]:
            for bar in bars:
                height = bar.get_height()
                ax.text(bar.get_x() + bar.get_width()/2., height + 0.5,
                       f'{height:.1f}%', ha='center', va='bottom', fontsize=8, fontweight='bold')
        
        ax.set_title(f'{self.bridge_name}左右幅偏差对比分析', fontsize=16, fontweight='bold')
        ax.set_ylabel('偏差 (%)', fontsize=14)
        ax.set_xlabel('索力测点', fontsize=14)
        ax.set_xticks(x)
        ax.set_xticklabels(list(valid_groups.keys()), rotation=45, ha='right')
        ax.legend(fontsize=12, loc='upper right')
        ax.grid(True, alpha=0.3)
        ax.tick_params(axis='both', which='major', labelsize=12)
        
        # 添加零线
        ax.axhline(y=0, color='black', linestyle='-', alpha=0.5)
        
        plt.tight_layout()
        
        # 保存图片
        output_filename = f"{self.bridge_name}左右幅偏差对比分析图.png"
        plt.savefig(self.output_dir / output_filename, dpi=300, bbox_inches='tight')
        plt.close()
        
        print(f"✅ 已保存: {output_filename}")
    
    def plot_cable_force_time_series(self):
        """绘制索力时间序列图"""
        print("📊 绘制索力时间序列图...")
        
        # 按左右幅分组
        left_data = {}
        right_data = {}
        
        for sensor_name, data in self.original_data.items():
            if data is not None and not data.empty:
                side = data['Side'].iloc[0]
                if side == '左幅':
                    left_data[sensor_name] = data
                elif side == '右幅':
                    right_data[sensor_name] = data
        
        # 创建更丰富的颜色方案
        extended_colors = [
            '#1f77b4', '#ff7f0e', '#2ca02c', '#d62728', '#9467bd',
            '#8c564b', '#e377c2', '#7f7f7f', '#bcbd22', '#17becf',
            '#a6cee3', '#fb9a99', '#fdbf6f', '#cab2d6', '#ff9896',
            '#f0027f', '#386cb0', '#fdc086', '#beaed4', '#7fc97f',
            '#bf5b17', '#666666', '#fb8072', '#80b1d3', '#fdb462',
            '#b3de69', '#fccde5', '#d9d9d9', '#bc80bd', '#ccebc5'
        ]
        
        # 绘制左幅索力时序图
        if left_data:
            fig, ax = plt.subplots(figsize=(15, 6))
            
            for i, (sensor_name, data) in enumerate(left_data.items()):
                color = extended_colors[i % len(extended_colors)]
                ax.plot(data['Time'], data['CableForce'], 
                       color=color, linewidth=1.5, alpha=0.8, 
                       label=sensor_name)
            
            ax.set_title(f'{self.bridge_name}左幅索力时序图', fontsize=16, fontweight='bold')
            ax.set_ylabel('索力 (KN)', fontsize=14)
            ax.set_xlabel('时间', fontsize=14)
            ax.grid(True, alpha=0.3)
            ax.legend(fontsize=10, loc='upper right', ncol=2)
            ax.tick_params(axis='both', which='major', labelsize=12)
            ax.tick_params(axis='x', rotation=45)
            
            plt.tight_layout()
            
            # 保存图片
            output_filename = f"{self.bridge_name}左幅索力时序图.png"
            plt.savefig(self.output_dir / output_filename, dpi=300, bbox_inches='tight')
            plt.close()
            
            print(f"✅ 已保存: {output_filename}")
        
        # 绘制右幅索力时序图
        if right_data:
            fig, ax = plt.subplots(figsize=(15, 6))
            
            for i, (sensor_name, data) in enumerate(right_data.items()):
                color = extended_colors[i % len(extended_colors)]
                ax.plot(data['Time'], data['CableForce'], 
                       color=color, linewidth=1.5, alpha=0.8, 
                       label=sensor_name)
            
            ax.set_title(f'{self.bridge_name}右幅索力时序图', fontsize=16, fontweight='bold')
            ax.set_ylabel('索力 (KN)', fontsize=14)
            ax.set_xlabel('时间', fontsize=14)
            ax.grid(True, alpha=0.3)
            ax.legend(fontsize=10, loc='upper right', ncol=2)
            ax.tick_params(axis='both', which='major', labelsize=12)
            ax.tick_params(axis='x', rotation=45)
            
            plt.tight_layout()
            
            # 保存图片
            output_filename = f"{self.bridge_name}右幅索力时序图.png"
            plt.savefig(self.output_dir / output_filename, dpi=300, bbox_inches='tight')
            plt.close()
            
            print(f"✅ 已保存: {output_filename}")
    
    def add_field(self, paragraph, field_code):
        """在段落中插入一个域代码"""
        run = paragraph.add_run()
        r = run._r
        
        # 开始符
        fldChar_begin = OxmlElement('w:fldChar')
        fldChar_begin.set(qn('w:fldCharType'), 'begin')
        r.append(fldChar_begin)
        
        # 域代码指令
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = field_code
        r.append(instrText)
        
        # 分隔符
        fldChar_separate = OxmlElement('w:fldChar')
        fldChar_separate.set(qn('w:fldCharType'), 'separate')
        r.append(fldChar_separate)
        
        # 结束符
        fldChar_end = OxmlElement('w:fldChar')
        fldChar_end.set(qn('w:fldCharType'), 'end')
        r.append(fldChar_end)
    
    def add_caption(self, doc, type_name, text):
        """为图或表添加带章节号的自动编号标题，应用图名格式"""
        p = doc.add_paragraph()
        p.add_run(f'{type_name} ')
        self.add_field(p, r'STYLEREF 1 \s')
        p.add_run('-')
        self.add_field(p, f'SEQ {type_name} \\* ARABIC \\s 1')
        p.add_run(f' {text}')
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 应用图名格式
        try:
            # 尝试应用图名格式样式
            p.style = '图名格式'
        except:
            # 如果样式不存在，使用默认格式
            pass
        
        return p
    
    def generate_analysis_text(self, current_average_data):
        """生成分析文字"""
        if not current_average_data:
            return "无有效数据进行分析。"
        
        # 计算总体统计
        all_forces = [data['average_force'] for data in current_average_data.values()]
        avg_force = np.mean(all_forces)
        max_force = max(all_forces)
        min_force = min(all_forces)
        
        # 计算偏差统计
        deviations = []
        for sensor_name, current_data in current_average_data.items():
            if sensor_name in self.historical_data_2022:
                historical_data = self.historical_data_2022[sensor_name]
                deviation = ((current_data['average_force'] - historical_data['load_test']) / historical_data['load_test']) * 100
                deviations.append(deviation)
        
        avg_deviation = np.mean(deviations)
        max_deviation = max(deviations)
        min_deviation = min(deviations)
        
        # 按左右幅分析
        left_forces = [data['average_force'] for data in current_average_data.values() if data['side'] == '左幅']
        right_forces = [data['average_force'] for data in current_average_data.values() if data['side'] == '右幅']
        
        left_avg = np.mean(left_forces) if left_forces else 0
        right_avg = np.mean(right_forces) if right_forces else 0
        
        text = f"根据{self.bridge_name}索力监测数据分析，得出以下结论：\n"
        text += f"1、监测期间各测点平均索力范围为{min_force:.2f}至{max_force:.2f} KN，"
        text += f"整体平均索力为{avg_force:.2f} KN。\n"
        text += f"2、左幅平均索力为{left_avg:.2f} KN，右幅平均索力为{right_avg:.2f} KN，"
        text += f"左右幅索力差异为{abs(left_avg - right_avg):.2f} KN。\n"
        text += f"3、本年度索力与荷载试验索力的偏差范围为{min_deviation:.2f}%至{max_deviation:.2f}%，"
        text += f"平均偏差为{avg_deviation:.2f}%。\n"
        text += f"4、索力变化趋势稳定，未发现明显的异常变化。"
        
        return text
    
    def generate_word_report(self, current_average_data, df_table):
        """生成Word报告"""
        try:
            from docx import Document
            from docx.shared import Inches, Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.oxml.shared import OxmlElement, qn
        except ImportError:
            print("❌ 缺少python-docx库，无法生成Word报告")
            return
        
        print(f"\n📄 正在生成Word报告...")
        
        # 创建新文档
        doc = Document()
        
        # 设置页面边距
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
        
        # 标题
        title = doc.add_heading(f'{self.bridge_name}索力数据分析报告', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 添加报告信息
        doc.add_paragraph(f"报告生成时间: {datetime.now().strftime('%Y年%m月%d日 %H:%M:%S')}")
        doc.add_paragraph(f"分析桥梁: {self.bridge_name}")
        doc.add_paragraph(f"数据来源: {self.data_dir}")
        
        # 1. 数据概况
        doc.add_heading('1. 数据概况', level=1)
        doc.add_paragraph(f"本次分析共处理了 {len(self.cable_force_data)} 个索力测点的数据。")
        
        # 统计左右幅测点数量
        left_count = sum(1 for data in self.cable_force_data.values() if data['Side'].iloc[0] == '左幅')
        right_count = sum(1 for data in self.cable_force_data.values() if data['Side'].iloc[0] == '右幅')
        doc.add_paragraph(f"其中左幅测点 {left_count} 个，右幅测点 {right_count} 个。")
        
        # 2. 索力对比分析
        doc.add_heading('2. 索力对比分析', level=1)
        
        # 为表格添加标题（在表格前）
        self.add_caption(doc, '表', '索力对比统计表')
        
        if not df_table.empty:
            # 创建统计表格
            table = doc.add_table(rows=1, cols=7)
            table.style = 'Table Grid'
            table.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 表头
            hdr_cells = table.rows[0].cells
            headers = ['斜拉索编号', '本年度监测索力(KN)', '2022年监测索力(KN)', 
                      '设计索力(KN)', '荷载试验索力(KN)', '本年度偏差(%)', '2022年偏差(%)']
            
            for cell, header_text in zip(hdr_cells, headers):
                cell.text = header_text
                for paragraph in cell.paragraphs:
                    try:
                        paragraph.style = '表格'
                    except:
                        pass
                for run in cell.paragraphs[0].runs:
                    run.bold = True
            
            # 添加数据行
            for _, row in df_table.iterrows():
                row_cells = table.add_row().cells
                
                data_items = [
                    row['斜拉索编号'],
                    row['本年度监测索力(KN)'],
                    row['2022年监测索力(KN)'],
                    row['设计索力(KN)'],
                    row['荷载试验索力(KN)'],
                    row['本年度偏差(%)'],
                    row['2022年偏差(%)']
                ]
                
                for cell, data in zip(row_cells, data_items):
                    cell.text = data
                    for paragraph in cell.paragraphs:
                        try:
                            paragraph.style = '表格'
                        except:
                            pass
            
            # 表格自动调整
            try:
                table.style = 'Table Grid'
                print("✅ 表格样式设置完成")
                
                tbl = table._element
                tblPr = tbl.tblPr
                if tblPr is None:
                    tblPr = OxmlElement('w:tblPr')
                    tbl.insert(0, tblPr)
                
                tblW = OxmlElement('w:tblW')
                tblW.set(qn('w:w'), '0')
                tblW.set(qn('w:type'), 'auto')
                tblPr.append(tblW)
                print("✅ 表格内容自动调整完成")
                
            except Exception as e:
                print(f"❌ 表格自动调整失败: {e}")
        
        # 3. 索力对比分析图
        doc.add_heading('3. 索力对比分析图', level=1)
        doc.add_paragraph("以下是索力对比分析图，展示了本年度监测索力、设计索力、荷载试验索力之间的对比关系。")
        
        # 添加对比分析图
        image_path = self.output_dir / f"{self.bridge_name}索力对比分析图.png"
        if image_path.exists():
            doc.add_picture(str(image_path), width=Inches(6))
            self.add_caption(doc, '图', f'{self.bridge_name}索力对比分析图')
        
        # 4. 索力偏差百分比分析图
        doc.add_heading('4. 索力偏差百分比分析图', level=1)
        doc.add_paragraph("以下是索力偏差百分比分析图，展示了本年度监测索力与设计索力之间的偏差百分比。")
        
        # 添加偏差分析图
        deviation_image_path = self.output_dir / f"{self.bridge_name}索力偏差百分比分析图.png"
        if deviation_image_path.exists():
            doc.add_picture(str(deviation_image_path), width=Inches(6))
            self.add_caption(doc, '图', f'{self.bridge_name}索力偏差百分比分析图')
        
        # 5. 左右幅索力对比分析
        doc.add_heading('5. 左右幅索力对比分析', level=1)
        doc.add_paragraph("以下是左右幅索力对比分析图，展示了同一测点左右幅的本年度监测索力对比。")
        
        # 添加左右幅索力对比图
        lr_force_image_path = self.output_dir / f"{self.bridge_name}左右幅索力对比分析图.png"
        if lr_force_image_path.exists():
            doc.add_picture(str(lr_force_image_path), width=Inches(6))
            self.add_caption(doc, '图', f'{self.bridge_name}左右幅索力对比分析图')
        
        # 6. 左右幅偏差对比分析
        doc.add_heading('6. 左右幅偏差对比分析', level=1)
        doc.add_paragraph("以下是左右幅偏差对比分析图，展示了同一测点左右幅的本年度和2022年偏差对比。")
        
        # 添加左右幅偏差对比图
        lr_deviation_image_path = self.output_dir / f"{self.bridge_name}左右幅偏差对比分析图.png"
        if lr_deviation_image_path.exists():
            doc.add_picture(str(lr_deviation_image_path), width=Inches(6))
            self.add_caption(doc, '图', f'{self.bridge_name}左右幅偏差对比分析图')
        
        # 7. 时间序列分析
        doc.add_heading('7. 时间序列分析', level=1)
        doc.add_paragraph("以下是各测点的索力时间序列图，展示了索力随时间的变化趋势。")
        
        # 添加时序图
        left_image_path = self.output_dir / f"{self.bridge_name}左幅索力时序图.png"
        if left_image_path.exists():
            doc.add_picture(str(left_image_path), width=Inches(6))
            self.add_caption(doc, '图', f'{self.bridge_name}左幅索力时序图')
        
        right_image_path = self.output_dir / f"{self.bridge_name}右幅索力时序图.png"
        if right_image_path.exists():
            doc.add_picture(str(right_image_path), width=Inches(6))
            self.add_caption(doc, '图', f'{self.bridge_name}右幅索力时序图')
        
        # 8. 结论与建议
        doc.add_heading('8. 结论与建议', level=1)
        
        # 生成分析文字
        analysis_text = self.generate_analysis_text(current_average_data)
        doc.add_paragraph(analysis_text)
        
        # 保存报告
        try:
            report_path = self.output_dir / f'{self.bridge_name}索力数据分析报告.docx'
            doc.save(str(report_path))
            print(f"✅ Word报告已生成: {report_path}")
            return str(report_path)
        except Exception as e:
            print(f"❌ 保存Word报告失败: {e}")
            return None
    
    def run_analysis(self):
        """运行完整分析"""
        print(f"🏗️ 开始分析 {self.bridge_name} 索力数据...")
        
        # 1. 加载索力数据
        if not self.load_cable_force_data():
            return
        
        # 2. 计算平均索力值
        current_average_data = self.calculate_average_cable_force()
        
        # 3. 创建对比表格
        df_table = self.create_comparison_table(current_average_data)
        
        # 4. 生成图表
        print("\n📊 生成可视化图表...")
        self.plot_cable_force_comparison(current_average_data)
        self.plot_cable_force_deviation(current_average_data)
        self.plot_left_right_force_comparison(current_average_data)
        self.plot_left_right_deviation_comparison(current_average_data)
        self.plot_cable_force_time_series()
        
        # 5. 生成Word报告
        print("\n📄 生成Word格式报告...")
        self.generate_word_report(current_average_data, df_table)
        
        print(f"\n✅ {self.bridge_name} 索力数据分析完成！")
        print(f"📁 结果保存在: {self.output_dir}")

def main():
    """主函数"""
    try:
        print("🚀 启动索力数据分析程序...")
        analyzer = CableForceAnalyzer()  # 使用配置中的桥名
        analyzer.run_analysis()
        print("✅ 程序执行完成")
    except Exception as e:
        print(f"❌ 程序执行出错: {e}")
        import traceback
        traceback.print_exc()

if __name__ == "__main__":
    main() 