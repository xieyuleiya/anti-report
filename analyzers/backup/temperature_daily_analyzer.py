#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
温度单日分析器
专门用于分析桥梁温度数据的单日详细分析
"""

import os
import pandas as pd
import matplotlib.pyplot as plt
import numpy as np
from pathlib import Path
from datetime import datetime, date
import warnings
warnings.filterwarnings('ignore')

# 设置中文字体
plt.rcParams['font.sans-serif'] = ['SimHei', 'Microsoft YaHei']
plt.rcParams['axes.unicode_minus'] = False

class TemperatureDailyAnalyzer:
    def __init__(self, bridge_name="白土北江特大桥"):
        # ==================== 配置区域 ====================
        self.bridge_name = bridge_name
        
        # 路径配置
        self.base_dir = Path("数据下载")
        self.data_dir = self.base_dir / bridge_name / "温度" / "数据"
        self.excel_path = Path("D:/desktop/全部通道.xlsx")
        self.output_dir = Path("wendu/分析结果")
        
        # 创建输出目录
        self.output_dir.mkdir(parents=True, exist_ok=True)
        
        # 数据存储
        self.excel_data = None
        self.file_mapping = {}
        self.temperature_data = {}
        self.grouped_data = {}
        
        print(f"🌡️ 温度单日分析器初始化完成")
        print(f"📁 数据目录: {self.data_dir}")
        print(f"📊 Excel文件: {self.excel_path}")
        print(f"📁 输出目录: {self.output_dir}")
    
    def load_excel_data(self):
        """加载Excel数据"""
        print(f"\n📊 正在加载Excel数据...")
        
        try:
            self.excel_data = pd.read_excel(self.excel_path)
            print(f"✅ Excel数据加载成功: {len(self.excel_data)} 行")
            
            # 筛选指定桥梁的数据
            filtered_data = self.excel_data[self.excel_data.iloc[:, 0] == self.bridge_name]
            print(f"✅ 筛选到 {self.bridge_name} 的数据: {len(filtered_data)} 行")
            
            self.build_file_mapping(filtered_data)
            
        except Exception as e:
            print(f"❌ Excel数据加载失败: {e}")
            return False
        
        return True
    
    def build_file_mapping(self, filtered_data):
        """构建文件映射关系"""
        print(f"\n🔍 正在构建文件映射关系...")
        
        # 获取温度文件夹中的所有txt文件
        if not self.data_dir.exists():
            print(f"❌ 数据目录不存在: {self.data_dir}")
            return False
        
        txt_files = list(self.data_dir.glob("*.txt"))
        print(f"📁 找到 {len(txt_files)} 个温度数据文件")
        
        # 为每个Excel行找到对应的数据文件
        for idx, row in filtered_data.iterrows():
            # 获取Excel行数据
            bridge_name = row.iloc[0]  # A列：桥名
            sensor_id = row.iloc[1]    # B列：通道ID
            sensor_name = row.iloc[2]  # C列：测点编号
            sensor_desc = row.iloc[3]  # D列：测点概况
            category = row.iloc[4]     # E列：测点所属种类
            location_large = row.iloc[5]  # F列：测点所属大位置
            location_small = row.iloc[6]  # G列：测点所属小位置
            
            # 只处理温度类型的数据
            if category != "温度":
                continue
            
            # 构建D+C的组合值用于匹配
            dc_combination = f"{sensor_desc}_{sensor_name}"
            
            # 查找匹配的文件
            matched_file = None
            for txt_file in txt_files:
                filename = txt_file.stem
                # 检查文件名是否包含D+C的组合或测点名称
                if dc_combination in filename or sensor_name in filename:
                    matched_file = txt_file
                    break
            
            if matched_file:
                # 存储文件映射和属性
                self.file_mapping[matched_file] = {
                    'bridge_name': bridge_name,
                    'sensor_id': sensor_id,
                    'sensor_name': sensor_name,
                    'sensor_desc': sensor_desc,
                    'category': category,
                    'location_large': location_large,
                    'location_small': location_small,
                    'group_key': f"{location_large}_{location_small}"
                }
                print(f"✅ 匹配成功: {sensor_name} -> {matched_file.name}")
            else:
                print(f"❌ 未找到匹配文件: {sensor_name} ({match_pattern})")
        
        print(f"📊 成功匹配 {len(self.file_mapping)} 个文件")
        return True
    
    def load_temperature_data(self):
        """加载温度数据"""
        print(f"\n📊 正在加载温度数据...")
        
        for file_path, attributes in self.file_mapping.items():
            sensor_name = attributes['sensor_name']
            print(f"📊 正在加载测点 {sensor_name} 的数据...")
            
            try:
                # 先打印原始文件前5行
                print(f"\n[调试] 文件: {file_path}")
                with open(file_path, encoding='utf-8', errors='ignore') as f:
                    for i in range(5):
                        line = f.readline()
                        if not line:
                            break
                        print(f"[原始] {line.strip()}")
                # 读取数据文件
                data = pd.read_csv(file_path, sep='\t', header=None, encoding='utf-8')
                print(f"[pandas] DataFrame前5行:")
                print(data.head())
                # 根据实际数据格式重新命名列
                # 第0列是时间，第2列是温度值
                if len(data.columns) >= 3:
                    data.columns = ['DateTime', 'Col1', 'Temperature'] + [f'Col{i}' for i in range(3, len(data.columns))]
                else:
                    data.columns = ['DateTime', 'Temperature']
                # 转换时间列 - 处理包含毫秒的格式
                data['Time'] = pd.to_datetime(data['DateTime'], format='%Y-%m-%d %H:%M:%S.%f', errors='coerce')
                # 如果毫秒格式解析失败，尝试其他格式
                if data['Time'].isna().all():
                    data['Time'] = pd.to_datetime(data['DateTime'], errors='coerce')
                # 删除多余的列，只保留时间和温度
                data = data[['Time', 'Temperature']]
                # 添加测点信息
                data['Sensor_Name'] = sensor_name
                data['Location_Large'] = attributes['location_large']
                data['Location_Small'] = attributes['location_small']
                data['Group_Key'] = attributes['group_key']
                # 存储数据
                self.temperature_data[sensor_name] = data
                print(f"✅ 成功加载测点 {sensor_name} 数据: {len(data)} 条记录")
                
            except Exception as e:
                print(f"❌ 加载文件 {file_path.name} 失败: {e}")
        
        # 按组别整理数据
        self.group_data_by_location()
    
    def group_data_by_location(self):
        """按F+G值分组整理数据"""
        print("\n正在按位置分组整理数据...")
        
        for sensor_name, data in self.temperature_data.items():
            if data is not None and not data.empty:
                group_key = data['Group_Key'].iloc[0]
                
                if group_key not in self.grouped_data:
                    self.grouped_data[group_key] = []
                
                self.grouped_data[group_key].append(data)
        
        print(f"📊 数据已按 {len(self.grouped_data)} 个组别整理")
        for group_key, group_data in self.grouped_data.items():
            print(f"  组别 {group_key}: {len(group_data)} 个测点")
    
    def get_available_dates(self):
        """获取所有可用日期"""
        all_dates = []
        for data in self.temperature_data.values():
            if data is not None and not data.empty:
                # 添加日期列
                data['Date'] = data['Time'].dt.date
                all_dates.extend(data['Date'].tolist())
        
        available_dates = sorted(list(set(all_dates)))
        return available_dates
    
    def extract_daily_data(self, selected_date):
        """提取指定日期的数据"""
        print(f"\n📅 正在提取 {selected_date} 的数据...")
        
        daily_data = {}
        
        for sensor_name, data in self.temperature_data.items():
            if data is not None and not data.empty:
                # 确保有Date列
                if 'Date' not in data.columns:
                    data['Date'] = data['Time'].dt.date
                
                # 筛选指定日期的数据
                daily_mask = data['Date'] == selected_date
                daily_subset = data[daily_mask].copy()
                
                if not daily_subset.empty:
                    daily_data[sensor_name] = daily_subset
                    print(f"✅ 测点 {sensor_name}: {len(daily_subset)} 条记录")
                else:
                    print(f"⚠️ 测点 {sensor_name}: 无数据")
        
        return daily_data
    
    def plot_daily_temperature_analysis(self, daily_data, selected_date):
        """绘制单日温度分析图"""
        if not daily_data:
            print("⚠️ 没有数据可绘制")
            return
        
        colors = ['blue', 'red', 'green', 'orange', 'purple', 'brown', 'pink', 'gray']
        
        # 按组别绘制
        for group_key, group_data in self.grouped_data.items():
            print(f"正在绘制组别 {group_key} 的单日温度图...")
            
            plt.figure(figsize=(15, 6))
            color_idx = 0
            
            for data in group_data:
                sensor_name = data['Sensor_Name'].iloc[0]
                if sensor_name in daily_data:
                    color = colors[color_idx % len(colors)]
                    
                    daily_subset = daily_data[sensor_name]
                    plt.plot(daily_subset['Time'], daily_subset['Temperature'], 
                            color=color, linewidth=1.5, alpha=0.8, 
                            label=sensor_name, marker='o', markersize=2)
                    color_idx += 1
            
            plt.title(f'{self.bridge_name} - {group_key}单日温度分析 ({selected_date})', fontsize=16, fontweight='bold')
            plt.ylabel('温度 (°C)', fontsize=14)
            plt.xlabel('时间', fontsize=14)
            plt.grid(True, alpha=0.3)
            plt.legend(fontsize=12)
            plt.tick_params(axis='both', which='major', labelsize=12)
            plt.xticks(rotation=45)
            plt.tight_layout()
            
            # 保存图片
            output_filename = f"{group_key}单日温度分析_{selected_date}.png"
            plt.savefig(self.output_dir / output_filename, dpi=300, bbox_inches='tight')
            plt.close()
            
            print(f"✅ 已保存: {output_filename}")
    
    def analyze_daily_statistics(self, daily_data):
        """分析单日统计数据"""
        if not daily_data:
            print("⚠️ 没有数据可分析")
            return {}
        
        daily_stats = {}
        
        for sensor_name, data in daily_data.items():
            if data is not None and not data.empty:
                # 确保温度列为数值类型
                data['Temperature'] = pd.to_numeric(data['Temperature'], errors='coerce')
                
                # 计算统计信息
                temp_stats = data['Temperature'].describe()
                
                daily_stats[sensor_name] = {
                    'temp_mean': temp_stats['mean'],
                    'temp_max': temp_stats['max'],
                    'temp_min': temp_stats['min'],
                    'temp_range': temp_stats['max'] - temp_stats['min'],
                    'record_count': len(data),
                    'time_range': f"{data['Time'].min().strftime('%H:%M')} - {data['Time'].max().strftime('%H:%M')}"
                }
        
        return daily_stats
    
    def generate_daily_report(self, daily_data, daily_stats, selected_date):
        """生成单日分析报告"""
        try:
            from docx import Document
            from docx.shared import Inches, Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError:
            print("❌ 缺少python-docx库，无法生成Word报告")
            return
        
        print(f"\n📄 正在生成单日分析报告...")
        
        # 创建新文档
        doc = Document()
        
        # 设置页面边距
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
        
        # 标题
        title = doc.add_heading(f'{self.bridge_name}温度单日分析报告', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 添加报告信息
        doc.add_paragraph(f"报告生成时间: {datetime.now().strftime('%Y年%m月%d日 %H:%M:%S')}")
        doc.add_paragraph(f"分析桥梁: {self.bridge_name}")
        doc.add_paragraph(f"分析日期: {selected_date}")
        doc.add_paragraph(f"数据来源: {self.data_dir}")
        
        # 1. 数据概况
        doc.add_heading('1. 数据概况', level=1)
        doc.add_paragraph(f"本次分析针对 {selected_date} 的温度数据，共处理了 {len(daily_data)} 个温度测点。")
        doc.add_paragraph(f"数据按位置分组，共分为 {len(self.grouped_data)} 个组别。")
        
        # 添加分组信息
        for group_key, group_data in self.grouped_data.items():
            doc.add_paragraph(f"• {group_key}: {len(group_data)} 个测点")
        
        # 2. 单日统计分析
        doc.add_heading('2. 单日统计分析', level=1)
        
        if daily_stats:
            # 创建统计表格
            table = doc.add_table(rows=1, cols=7)
            table.style = 'Table Grid'
            
            # 表头
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = '测点名称'
            hdr_cells[1].text = '平均温度(°C)'
            hdr_cells[2].text = '最高温度(°C)'
            hdr_cells[3].text = '最低温度(°C)'
            hdr_cells[4].text = '温度差值(°C)'
            hdr_cells[5].text = '记录数量'
            
            # 添加数据行
            for sensor_name, stats in daily_stats.items():
                row_cells = table.add_row().cells
                row_cells[0].text = sensor_name
                row_cells[1].text = f"{stats['temp_mean']:.2f}"
                row_cells[2].text = f"{stats['temp_max']:.2f}"
                row_cells[3].text = f"{stats['temp_min']:.2f}"
                row_cells[4].text = f"{stats['temp_range']:.2f}"
                row_cells[5].text = str(stats['record_count'])
        
        # 3. 单日温度变化分析
        doc.add_heading('3. 单日温度变化分析', level=1)
        doc.add_paragraph("以下是各测点组的单日温度变化图，展示了温度在一天内的变化趋势。")
        
        # 添加图片
        for group_key in self.grouped_data.keys():
            image_path = self.output_dir / f"{group_key}单日温度分析_{selected_date}.png"
            if image_path.exists():
                doc.add_heading(f'{group_key}单日温度分析', level=2)
                doc.add_picture(str(image_path), width=Inches(6))
                doc.add_paragraph(f"图3-{list(self.grouped_data.keys()).index(group_key)+1}: {group_key}单日温度分析图")
        
        # 4. 结论
        doc.add_heading('4. 结论', level=1)
        doc.add_paragraph("基于以上分析，可以得出以下结论：")
        
        if daily_stats:
            # 计算总体统计
            all_temps = []
            for stats in daily_stats.values():
                all_temps.extend([stats['temp_mean']])
            
            overall_mean = np.mean(all_temps)
            overall_max = max([stats['temp_max'] for stats in daily_stats.values()])
            overall_min = min([stats['temp_min'] for stats in daily_stats.values()])
            
            doc.add_paragraph(f"• {selected_date} 所有测点的平均温度为 {overall_mean:.2f}°C")
            doc.add_paragraph(f"• 当日最高温度出现在 {overall_max:.2f}°C")
            doc.add_paragraph(f"• 当日最低温度出现在 {overall_min:.2f}°C")
            doc.add_paragraph(f"• 当日温度变化范围为 {overall_max - overall_min:.2f}°C")
        
        # 保存报告
        report_path = self.output_dir / f'{self.bridge_name}温度单日分析报告_{selected_date}.docx'
        doc.save(str(report_path))
        print(f"✅ 单日分析报告已保存: {report_path}")
    
    def run_analysis(self, selected_date=None):
        """运行单日分析"""
        print(f"🌡️ 开始分析 {self.bridge_name} 温度单日数据...")
        
        # 加载Excel数据
        if not self.load_excel_data():
            return
        
        # 加载温度数据
        self.load_temperature_data()
        
        # 获取可用日期
        available_dates = self.get_available_dates()
        if not available_dates:
            print("❌ 没有可用的日期数据")
            return
        
        print(f"📅 可用日期范围: {available_dates[0]} 到 {available_dates[-1]}")
        print(f"📅 共有 {len(available_dates)} 天的数据")
        
        # 如果没有指定日期，使用第一个可用日期
        if selected_date is None:
            selected_date = available_dates[0]
            print(f"🔄 使用默认日期: {selected_date}")
        
        # 提取单日数据
        daily_data = self.extract_daily_data(selected_date)
        if not daily_data:
            print(f"❌ 未找到 {selected_date} 的数据")
            return
        
        # 绘制单日分析图
        self.plot_daily_temperature_analysis(daily_data, selected_date)
        
        # 分析单日统计
        daily_stats = self.analyze_daily_statistics(daily_data)
        
        print(f"\n📊 {selected_date} 单日分析结果:")
        for sensor_name, stats in daily_stats.items():
            print(f"  测点 {sensor_name}:")
            print(f"    平均温度: {stats['temp_mean']:.2f}°C")
            print(f"    最高温度: {stats['temp_max']:.2f}°C")
            print(f"    最低温度: {stats['temp_min']:.2f}°C")
            print(f"    温度差值: {stats['temp_range']:.2f}°C")
            print(f"    记录数量: {stats['record_count']}")
        
        # 生成单日分析报告
        self.generate_daily_report(daily_data, daily_stats, selected_date)
        
        print(f"\n✅ {self.bridge_name} 温度单日分析完成！")
        print(f"📁 结果保存在: {self.output_dir}")

def main():
    """主函数"""
    print("🚀 温度单日分析器启动...")
    
    # 可以在这里修改桥名和日期
    bridge_name = "白土北江特大桥"
    selected_date = date(2024, 7, 1)  # 可以修改为其他日期
    
    print(f"📋 分析目标桥梁: {bridge_name}")
    print(f"📅 分析日期: {selected_date}")
    
    try:
        analyzer = TemperatureDailyAnalyzer(bridge_name)
        print("✅ 分析器初始化成功")
        analyzer.run_analysis(selected_date)
    except Exception as e:
        print(f"❌ 程序运行出错: {e}")
        import traceback
        traceback.print_exc()

if __name__ == "__main__":
    main() 